# -*- coding: utf-8 -*-
"""Generate Cairo Bags Frontend API Integration Architecture Word report on Desktop."""
import io
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")
OUTPUT = os.path.join(DESKTOP, "Cairo_Bags_Frontend_API_Integration_Architecture_v2.docx")

HEADER = "1A1A1A"
GOLD = "C9A962"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def add_table(doc, headers, rows, header_color=HEADER):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph()


def make_global_data_flow_diagram():
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.5)
    ax.axis("off")

    def box(x, y, w, h, text, color="#FAFAFA", fs=8):
        patch = FancyBboxPatch(
            (x - w / 2, y - h / 2), w, h,
            boxstyle="round,pad=0.02", facecolor=color, edgecolor="#0A0A0A", linewidth=1.2)
        ax.add_patch(patch)
        ax.text(x, y, text, ha="center", va="center", fontsize=fs, fontweight="bold", color="#0A0A0A")

    box(2.5, 4.5, 3.2, 0.7, "Contexts\n(Auth, Cart, Notifications)", "#FFF8E7", fs=8)
    box(7.5, 4.5, 3.0, 0.7, "Pages / Hooks\n(future UI)", "#FAFAFA", fs=8)
    box(5.0, 3.0, 3.4, 0.7, "Domain Services\n(authService, cartService, ...)", "#F0F0F0", fs=8)
    box(5.0, 1.6, 3.0, 0.7, "axiosInstance\n(Single Source of Truth)", "#E8E8E8", fs=8)
    box(2.5, 0.4, 2.8, 0.65, "REST API\nCairo Bags Backend", "#C9A962", fs=8)
    box(7.5, 0.4, 2.8, 0.65, "SignalR Hub\n/hubs/notifications", "#E8D5A3", fs=8)
    box(8.8, 3.0, 1.6, 0.55, "localStorage\ntoken, user", "#FFFFFF", fs=7)

    arrows = [
        ((2.5, 4.1), (4.2, 3.35)), ((7.5, 4.1), (5.8, 3.35)),
        ((5.0, 2.65), (5.0, 1.95)), ((4.2, 1.6), (3.0, 0.75)),
        ((5.8, 1.6), (7.0, 0.75)), ((6.2, 1.6), (8.2, 2.7)),
    ]
    for (x1, y1), (x2, y2) in arrows:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color="#C9A962", lw=1.3))

    ax.set_title("Cairo Bags — Frontend API Data Flow", fontsize=13, fontweight="bold", color="#0A0A0A")
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def make_auth_flow_diagram():
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    fig, ax = plt.subplots(figsize=(10, 3.8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.8)
    ax.axis("off")

    steps = [
        (1.0, "Login /\nRegister"),
        (3.0, "Save token +\nrefreshToken + user"),
        (5.0, "Merge Guest\nCart"),
        (7.0, "Connect\nSignalR"),
        (9.0, "Authenticated\nSession"),
    ]
    for x, label in steps:
        patch = FancyBboxPatch(
            (x - 0.85, 1.2), 1.7, 1.0,
            boxstyle="round,pad=0.03", facecolor="#FFF8E7", edgecolor="#0A0A0A", linewidth=1.2)
        ax.add_patch(patch)
        ax.text(x, 1.7, label, ha="center", va="center", fontsize=8, fontweight="bold")

    for x1, x2 in [(1.85, 2.15), (3.85, 4.15), (5.85, 6.15), (7.85, 8.15)]:
        ax.annotate("", xy=(x2, 1.7), xytext=(x1, 1.7),
                    arrowprops=dict(arrowstyle="->", color="#0A0A0A", lw=1.2))

    ax.text(5, 3.2, "401 → refresh-token (queue) → retry or logout", ha="center", fontsize=9, color="#555555")
    ax.set_title("Authentication & Session Lifecycle", fontsize=13, fontweight="bold", color="#0A0A0A")

    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def make_order_payment_flow_diagram():
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    fig, ax = plt.subplots(figsize=(10, 4.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.2)
    ax.axis("off")

    steps = [
        (1.0, "Cart"),
        (2.8, "Checkout"),
        (4.6, "Awaiting\nPayment"),
        (6.4, "Proof\nUpload"),
        (8.2, "Admin\nReview"),
    ]
    for x, label in steps:
        patch = FancyBboxPatch(
            (x - 0.75, 1.5), 1.5, 0.9,
            boxstyle="round,pad=0.02", facecolor="#FAFAFA", edgecolor="#0A0A0A", linewidth=1.1)
        ax.add_patch(patch)
        ax.text(x, 1.95, label, ha="center", va="center", fontsize=8, fontweight="bold")

    admin_steps = [(2.8, "Processing"), (4.6, "Shipped"), (6.4, "Delivered")]
    for x, label in admin_steps:
        patch = FancyBboxPatch(
            (x - 0.7, 0.35), 1.4, 0.65,
            boxstyle="round,pad=0.02", facecolor="#C9A962", edgecolor="#0A0A0A", linewidth=1.1)
        ax.add_patch(patch)
        ax.text(x, 0.67, label, ha="center", va="center", fontsize=7, fontweight="bold")

    for x1, x2 in [(1.75, 2.05), (3.55, 3.85), (5.35, 5.65), (7.15, 7.45)]:
        ax.annotate("", xy=(x2, 1.95), xytext=(x1, 1.95),
                    arrowprops=dict(arrowstyle="->", color="#0A0A0A", lw=1.1))

    ax.annotate("", xy=(2.8, 1.0), xytext=(8.2, 1.5),
                arrowprops=dict(arrowstyle="->", color="#C9A962", lw=1.1, connectionstyle="arc3,rad=-0.3"))
    ax.text(5.5, 0.1, "After payment approved → admin order transitions", ha="center", fontsize=8, color="#555555")
    ax.set_title("Order & Payment Integration Flow", fontsize=13, fontweight="bold", color="#0A0A0A")

    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def build_document():
    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Cairo Bags")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Frontend API Integration Architecture")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x62)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Service Layer Design — Axios Single Source of Truth — No UI Code")
    run.font.size = Pt(12)
    run.font.italic = True

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub3.add_run(datetime.now().strftime("%B %Y"))
    run.font.size = Pt(11)

    doc.add_page_break()

    # Executive Summary
    add_heading(doc, "Executive Summary", 1)
    add_para(doc,
             "This document defines the complete API integration architecture for the Cairo Bags React frontend. "
             "All HTTP communication flows through a single Axios instance (master configuration file). "
             "No additional Axios clients, no hardcoded base URLs outside environment configuration.")
    add_table(doc, ["Principle", "Decision"], [
        ["HTTP client", "Single axiosInstance — mandatory foundation"],
        ["Base URL", "VITE_API_BASE_URL only (e.g. http://localhost:5073)"],
        ["Endpoint truth", "Swagger /swagger/v1/swagger.json is FINAL authority"],
        ["Auth", "JWT Bearer + refresh token (body + HttpOnly cookie)"],
        ["Guest cart", "X-Session-Id header on cart/recommendation calls"],
        ["Realtime", "SignalR /hubs/notifications (parallel to Axios)"],
        ["Scope", "Architecture & integration design only — no React UI"],
    ])
    add_para(doc,
             "All endpoints below were verified by scanning backend Controllers and cross-checked "
             "against live Swagger (74 routes). Any Swagger-only route must be included even if "
             "absent from earlier architecture drafts.")

    doc.add_page_break()

    # Axios Foundation
    add_heading(doc, "0. Axios Foundation Analysis", 1)
    add_para(doc, "The attached axios configuration is the SINGLE SOURCE OF TRUTH. Preserve:", bold=True)
    add_bullet(doc, "JWT Authorization header from localStorage token")
    add_bullet(doc, "Refresh token flow with isRefreshing + failedQueue")
    add_bullet(doc, "Request and response interceptors")
    add_bullet(doc, "withCredentials: true for cookie-based refresh")
    add_bullet(doc, "localStorage strategy: token + user (extend with refreshToken)")

    add_heading(doc, "Backend Alignment Required", 2)
    add_table(doc, ["Axios file (current)", "Cairo Bags backend (actual)"], [
        ["baseURL: https://localhost:7250", "VITE_API_BASE_URL → http://localhost:5073 (dev)"],
        ["/api/Accounts/RefreshToken", "POST /api/Account/refresh-token"],
        ["/api/Accounts/Login", "POST /api/Account/LogIn"],
        ["Refresh body: {}", "{ refreshToken } + withCredentials cookie"],
    ])
    add_para(doc, "Update paths inside the same axios file — do NOT create a second instance.")

    doc.add_page_break()

    # Folder structure
    add_heading(doc, "1. Recommended Folder Structure", 1)
    add_code_block(doc, """src/
├── api/
│   ├── axios.js                 # re-export master instance
│   ├── index.js                 # barrel export
│   ├── constants/
│   │   ├── endpoints.js
│   │   ├── orderStatus.js
│   │   ├── paymentStatus.js
│   │   └── paymentMethods.js
│   ├── utils/
│   │   ├── normalizeError.js
│   │   ├── sessionId.js
│   │   ├── authStorage.js
│   │   └── multipart.js
│   ├── auth/authService.js
│   ├── categories/categoryService.js
│   ├── products/
│   │   ├── productService.js
│   │   └── productImageService.js
│   ├── inventory/inventoryService.js
│   ├── cart/cartService.js
│   ├── checkout/checkoutService.js
│   ├── orders/
│   │   ├── orderService.js
│   │   └── adminOrderService.js
│   ├── payments/
│   │   ├── paymentService.js
│   │   └── adminPaymentService.js
│   ├── recommendations/recommendationService.js
│   ├── notifications/
│   │   ├── notificationService.js
│   │   └── notificationHub.js
│   ├── files/fileService.js
│   └── system/systemSettingsService.js
├── context/                     # consumes services only
└── hooks/""")

    add_para(doc, "Rule: Only api/axios.js imports the master instance. Every service imports from ../axios.js.")

    doc.add_page_break()

    # Service organization
    add_heading(doc, "2. API Service Organization", 1)

    add_heading(doc, "authService.js", 2)
    add_para(doc, "Responsibility: Identity, session lifecycle, profile.")
    add_table(doc, ["Method", "HTTP", "Endpoint", "Auth"], [
        ["register", "POST", "/api/Account/register", "Anonymous"],
        ["login", "POST", "/api/Account/LogIn", "Anonymous"],
        ["refreshToken", "POST", "/api/Account/refresh-token", "Anonymous + refresh cookie/body"],
        ["logout", "POST", "/api/Account/LogOut", "Required"],
        ["getMe", "GET", "/api/Account/Me", "Required"],
        ["updateMe", "PUT", "/api/Account/Me", "Required"],
        ["updateUsername", "PUT", "/api/Account/update-username", "Required"],
        ["googleSignIn", "POST", "/api/Account/sign-in-google", "Anonymous"],
        ["forgotPasswordRequest", "POST", "/api/Account/forgot-password/request-code", "Anonymous"],
        ["forgotPasswordComplete", "POST", "/api/Account/forgot-password/complete", "Anonymous"],
        ["changePassword", "POST", "/api/Account/change-password", "Required"],
        ["setPassword", "POST", "/api/Account/set-password", "Required (Google users)"],
        ["markFirstLoginDone", "POST", "/api/Account/mark-first-login-done", "Required"],
        ["createAdmin", "POST", "/api/Account/create-admin", "Anonymous* / Admin if exists"],
    ])
    add_para(doc, "updateUsername body: { username }. setPassword body: { newPassword }. "
             "markFirstLoginDone: no body.")

    add_heading(doc, "categoryService.js", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["getCategories", "GET", "/api/categories"],
        ["getCategoryTree", "GET", "/api/categories/tree"],
        ["getCategoryById", "GET", "/api/categories/{id}"],
        ["createCategory", "POST", "/api/admin/categories"],
        ["updateCategory", "PUT", "/api/admin/categories/{id}"],
        ["deleteCategory", "DELETE", "/api/admin/categories/{id}"],
    ])

    add_heading(doc, "productService.js", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["getProducts", "GET", "/api/products"],
        ["getFeatured", "GET", "/api/products/featured"],
        ["getNewArrivals", "GET", "/api/products/new-arrivals"],
        ["searchProducts", "GET", "/api/products/search"],
        ["getProductById", "GET", "/api/products/{id}"],
        ["createProduct", "POST", "/api/admin/products"],
        ["updateProduct", "PUT", "/api/admin/products/{id}"],
        ["deleteProduct", "DELETE", "/api/admin/products/{id}"],
    ])

    add_heading(doc, "productImageService.js", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["getProductImages", "GET", "/api/products/{productId}/images"],
        ["uploadImage", "POST", "/api/admin/products/{productId}/images"],
        ["uploadVariantImage", "POST", "/api/admin/products/{productId}/images/variant/{variantId}"],
        ["setPrimary", "PUT", "/api/admin/products/{productId}/images/{imageId}/primary"],
        ["reorder", "PUT", "/api/admin/products/{productId}/images/reorder"],
        ["deleteImage", "DELETE", "/api/admin/products/{productId}/images/{imageId}"],
    ])
    add_para(doc, "Uploads use multipart/form-data via FormData; omit Content-Type header on those requests.")

    doc.add_page_break()

    add_heading(doc, "inventoryService.js", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["getVariantStatus", "GET", "/api/inventory/{variantId}/status"],
        ["listInventory", "GET", "/api/admin/inventory"],
        ["getLowStock", "GET", "/api/admin/inventory/low-stock"],
        ["getVariantInventory", "GET", "/api/admin/inventory/{variantId}"],
        ["getMovements", "GET", "/api/admin/inventory/{variantId}/movements"],
        ["adjustStock", "POST", "/api/admin/inventory/{variantId}/adjust"],
        ["reserve", "POST", "/api/admin/inventory/{variantId}/reserve"],
        ["release", "POST", "/api/admin/inventory/{variantId}/release"],
    ])

    add_heading(doc, "cartService.js", 2)
    add_para(doc, "Responsibility: Guest + authenticated cart. Guest requires X-Session-Id header.")
    add_table(doc, ["Method", "HTTP", "Endpoint", "Auth"], [
        ["getCart", "GET", "/api/cart", "Optional"],
        ["addItem", "POST", "/api/cart/items", "Optional"],
        ["updateItem", "PUT", "/api/cart/items/{variantId}", "Optional"],
        ["removeItem", "DELETE", "/api/cart/items/{variantId}", "Optional"],
        ["clearCart", "DELETE", "/api/cart", "Optional"],
        ["mergeCart", "POST", "/api/cart/merge", "Required"],
    ])

    add_heading(doc, "checkoutService.js", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["checkout", "POST", "/api/checkout"],
    ])
    add_para(doc, "Body: shippingAddressId, paymentMethod (enum 1-6), couponCode?, notes?")

    add_heading(doc, "orderService.js (customer)", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["getMyOrders", "GET", "/api/orders"],
        ["getOrderById", "GET", "/api/orders/{id}"],
        ["cancelOrder", "POST", "/api/orders/{id}/cancel"],
    ])

    add_heading(doc, "adminOrderService.js", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["listOrders", "GET", "/api/admin/orders"],
        ["getOrder", "GET", "/api/admin/orders/{id}"],
        ["markProcessing", "POST", "/api/admin/orders/{id}/processing"],
        ["markShipped", "POST", "/api/admin/orders/{id}/shipped"],
        ["markDelivered", "POST", "/api/admin/orders/{id}/delivered"],
        ["cancelOrder", "POST", "/api/admin/orders/{id}/cancel"],
        ["refundOrder", "POST", "/api/admin/orders/{id}/refund"],
    ])

    doc.add_page_break()

    add_heading(doc, "paymentService.js (customer)", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["submitProof", "POST", "/api/payments/{orderId}/proof"],
        ["getPaymentByOrder", "GET", "/api/payments/{orderId}"],
    ])
    add_para(doc, "submitProof: multipart — SenderName, SenderPhone, TransactionReference, ProofFiles[]")

    add_heading(doc, "adminPaymentService.js", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["getPending", "GET", "/api/admin/payments/pending"],
        ["getPaymentById", "GET", "/api/admin/payments/{paymentId}"],
        ["approve", "POST", "/api/admin/payments/{paymentId}/approve"],
        ["reject", "POST", "/api/admin/payments/{paymentId}/reject"],
    ])

    add_heading(doc, "recommendationService.js", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["getTrending", "GET", "/api/recommendations/trending"],
        ["getRecentlyViewed", "GET", "/api/recommendations/recently-viewed"],
        ["getSimilar", "GET", "/api/recommendations/similar/{productId}"],
        ["getFrequentlyBoughtTogether", "GET", "/api/recommendations/frequently-bought-together/{productId}"],
    ])

    add_heading(doc, "notificationService.js + notificationHub.js", 2)
    add_table(doc, ["Transport", "Method/Event", "Endpoint/Hub"], [
        ["REST", "GET", "/api/Notifications?page&pageSize"],
        ["REST", "POST", "/api/Notifications/read/{id}"],
        ["REST", "POST", "/api/Notifications/read-all"],
        ["REST", "GET", "/api/Notifications/unread-count"],
        ["SignalR", "ReceiveNotification", "/hubs/notifications"],
        ["SignalR", "UnreadCountUpdated", "/hubs/notifications"],
    ])

    add_heading(doc, "fileService.js", 2)
    add_para(doc, "Responsibility: Generic authenticated file upload. Returns relative url for use in other APIs.")
    add_table(doc, ["Method", "HTTP", "Endpoint", "Auth"], [
        ["uploadFile", "POST", "/api/File/Upload", "Required"],
    ])
    add_para(doc, "Request: multipart/form-data with field file. Response: { url: \"/FileStorage/{guid}.ext\" }")
    add_para(doc, "Allowed: jpg, jpeg, png, gif, webp, webm, mp3, wav. Max 10 MB.")

    add_heading(doc, "systemSettingsService.js", 2)
    add_table(doc, ["Method", "HTTP", "Endpoint"], [
        ["getSettings", "GET", "/api/SystemSettings"],
        ["updateSettings", "PUT", "/api/SystemSettings"],
    ])

    doc.add_page_break()

    # Image upload flow
    add_heading(doc, "2.1 Image Upload Flow (Verified)", 1)
    add_para(doc,
             "Cairo Bags supports TWO image patterns. Frontend must implement both where applicable:")
    add_heading(doc, "Pattern A — Generic File Upload → imageUrl", 2)
    add_bullet(doc, "Step 1: POST /api/File/Upload (multipart file) → receive { url }")
    add_bullet(doc, "Step 2: Prefix url with API host for display: VITE_API_BASE_URL + url")
    add_bullet(doc, "Step 3: Pass url string into downstream JSON bodies as imageUrl / ImageUrl / ProfileImageUrl")
    add_table(doc, ["Target API", "Field", "When"], [
        ["POST/PUT /api/admin/categories", "imageUrl", "Category create/update"],
        ["POST/PUT /api/admin/products", "images[].imageUrl", "Product create/update with ProductImageInputDto"],
        ["PUT /api/Account/Me", "profileImageUrl", "Customer profile avatar"],
    ])
    add_heading(doc, "Pattern B — Direct Product Image Upload (multipart)", 2)
    add_bullet(doc, "POST /api/admin/products/{productId}/images — file + form altText/isPrimary")
    add_bullet(doc, "POST /api/admin/products/{productId}/images/variant/{variantId} — variant-specific image")
    add_bullet(doc, "Use when uploading to an existing product without pre-uploading via File/Upload")
    add_heading(doc, "Pattern C — Payment Proof (customer)", 2)
    add_bullet(doc, "POST /api/payments/{orderId}/proof — multipart ProofFiles[] (not File/Upload)")
    add_para(doc, "Recommended admin catalog flow: Pattern A for category hero images and product JSON; "
             "Pattern B for gallery management on product edit screen.")

    doc.add_page_break()

    # Auth integration
    add_heading(doc, "3. Authentication Integration Plan", 1)
    doc.add_picture(io.BytesIO(make_auth_flow_diagram()), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "localStorage Strategy", 2)
    add_table(doc, ["Key", "Content"], [
        ["token", "JWT access token"],
        ["refreshToken", "Opaque refresh token (body fallback)"],
        ["user", "Serialized user profile + roles from login/Me"],
        ["guestSessionId", "UUID for X-Session-Id header"],
    ])

    add_heading(doc, "JWT Role Detection", 2)
    add_table(doc, ["Role", "Detection"], [
        ["Guest", "No token in localStorage"],
        ["Customer", "token present + role includes Customer"],
        ["Admin", "token present + role includes Admin"],
    ])
    add_para(doc, "JWT claims: sub (user id), email, role. Confirm with GET /api/Account/Me on app bootstrap.")

    add_heading(doc, "AuthContext Responsibilities", 2)
    add_bullet(doc, "Call authService only — never raw axios")
    add_bullet(doc, "On login/register: persist tokens + user, then cartService.mergeCart(sessionId)")
    add_bullet(doc, "On boot: if token exists → getMe() to refresh profile")
    add_bullet(doc, "On logout: authService.logout() + clear storage + disconnect SignalR")

    doc.add_page_break()

    # Route protection
    add_heading(doc, "4. Route Protection Strategy", 1)
    add_table(doc, ["Zone", "Examples", "Guard"], [
        ["Public", "Home, shop, product, login, register", "None"],
        ["Customer", "Checkout, my orders, payment proof, profile", "isAuthenticated"],
        ["Admin", "/admin/* catalog, inventory, orders, payments", "isAuthenticated && isAdmin"],
    ])
    add_para(doc, "Router guards: PublicRoute, CustomerRoute (redirect to /login), AdminRoute (redirect to / or /403).")
    add_para(doc, "Backend enforces [Authorize] independently — frontend guards are UX only.")

    # Error handling
    add_heading(doc, "5. Error Handling Strategy", 1)
    add_para(doc, "All services use normalizeError(error) → unified shape:")
    add_code_block(doc, """{
  status: number | null,
  code: string | null,       // e.g. "order_not_found"
  message: string,
  details: object | null,
  isNetworkError: boolean,
  isAuthError: boolean
}""")

    add_table(doc, ["Status", "Meaning", "Frontend Action"], [
        ["400", "Validation / business rule", "Show field errors from code + message"],
        ["401", "Unauthorized", "Interceptor refresh; if still 401 → logout + redirect"],
        ["403", "Forbidden", "Show permission error; admin routes redirect"],
        ["404", "Not found", "Empty state / not found page"],
        ["409", "Conflict (stock, order state)", "Show message; checkout inventory conflict"],
        ["429", "Rate limit (auth)", "Too many attempts message"],
        ["500", "Server error", "Generic message; log detail in dev"],
    ])

    doc.add_page_break()

    # Notifications
    add_heading(doc, "6. Notification Integration Strategy", 1)
    add_para(doc, "NotificationContext orchestrates REST + SignalR — never raw axios in UI.")
    add_table(doc, ["State", "Source"], [
        ["notifications[]", "GET /api/Notifications"],
        ["unreadCount", "GET unread-count + hub UnreadCountUpdated"],
        ["connect()", "On login — hub with JWT accessTokenFactory"],
        ["disconnect()", "On logout"],
        ["markRead(id)", "POST read/{id} + optimistic update"],
    ])
    add_para(doc, "Push event ReceiveNotification payload: id, type, targetType, referenceId, deepLink, title, message, createdAt, isRead")
    add_para(doc, "Types: order_placed, payment_confirmed, payment_rejected, order_shipped, order_delivered, etc.")

    # Cart sync
    add_heading(doc, "7. Cart Synchronization Strategy", 1)
    add_bullet(doc, "Guest: no JWT; all cart calls send X-Session-Id header (localStorage guestSessionId)")
    add_bullet(doc, "Authenticated: JWT identifies user; session header optional")
    add_bullet(doc, "After login/register: always POST /api/cart/merge with guest sessionId before refreshing cart UI")
    add_bullet(doc, "Surface CartItemDto.stockChanged and maxAllowedQuantity in CartContext")
    add_bullet(doc, "Checkout requires auth; cart must be user-owned post-merge")

    doc.add_page_break()

    # Order & Payment flow
    add_heading(doc, "8. Order & Payment Integration Flow", 1)
    doc.add_picture(io.BytesIO(make_order_payment_flow_diagram()), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Backend Status Enums", 2)
    add_para(doc, "OrderStatus: Pending → AwaitingPayment → PaymentProofSubmitted → PaymentUnderReview → "
             "PaymentConfirmed → Processing → Shipped → Delivered → Completed | Cancelled | Refunded")
    add_para(doc, "PaymentStatus: Pending → ProofSubmitted → UnderReview → Confirmed | Rejected | Refunded")

    add_heading(doc, "Screen → API Mapping", 2)
    add_table(doc, ["Frontend Step", "API Calls"], [
        ["Cart", "GET/POST/PUT/DELETE /api/cart*"],
        ["Checkout", "POST /api/checkout"],
        ["Awaiting Payment", "GET /api/payments/{orderId}"],
        ["Proof Upload", "POST /api/payments/{orderId}/proof (multipart)"],
        ["Payment Under Review", "GET /api/orders/{id} + GET /api/payments/{orderId}"],
        ["Admin Pending Payments", "GET /api/admin/payments/pending"],
        ["Admin Approve/Reject", "POST /api/admin/payments/{id}/approve|reject"],
        ["Processing", "POST /api/admin/orders/{id}/processing"],
        ["Shipped", "POST /api/admin/orders/{id}/shipped"],
        ["Delivered", "POST /api/admin/orders/{id}/delivered"],
        ["My Orders", "GET /api/orders"],
        ["Order Detail / Cancel", "GET /api/orders/{id}, POST .../cancel"],
    ])

    add_heading(doc, "Payment Method Enum", 2)
    add_table(doc, ["Value", "Method"], [
        ["1", "CashOnDelivery"],
        ["2", "InstaPay"],
        ["3", "VodafoneCash"],
        ["4", "OrangeCash"],
        ["5", "EtisalatCash"],
        ["6", "WEPay"],
    ])

    doc.add_page_break()

    # Data flow + service principles
    add_heading(doc, "9. Global Data Flow", 1)
    doc.add_picture(io.BytesIO(make_global_data_flow_diagram()), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "10. Service Layer Design Principles", 1)
    principles = [
        "One function = one endpoint (thin wrappers)",
        "Return response.data from services (not full axios response)",
        "No UI logic in services",
        "Multipart: FormData + delete Content-Type on upload requests only",
        "Session-aware: cartService and recommendationService attach X-Session-Id",
        "Separate admin services (adminOrderService, adminPaymentService)",
        "Barrel export via api/index.js for contexts",
    ]
    for p in principles:
        add_bullet(doc, p)

    add_heading(doc, "11. Environment Configuration", 1)
    add_code_block(doc, """# .env
VITE_API_BASE_URL=http://localhost:5073

# axios.js
baseURL: import.meta.env.VITE_API_BASE_URL""")

    add_heading(doc, "12. Pre-Implementation Checklist", 1)
    checklist = [
        "Align axios interceptor paths to /api/Account/*",
        "Set VITE_API_BASE_URL — no hardcoded URLs elsewhere",
        "Store refreshToken on login; send in refresh body",
        "Implement guestSessionId + X-Session-Id for cart/recommendations",
        "mergeCart immediately after successful auth",
        "normalizeError utility used by all services",
        "NotificationContext: REST bootstrap + SignalR realtime",
        "Admin routes check JWT role === Admin",
        "Multipart helpers for payment proof + product images",
        "fileService.js for POST /api/File/Upload",
        "Image flow: File/Upload → url → imageUrl in Categories/Products/Profile",
        "Auth: update-username, set-password, mark-first-login-done wired in authService",
    ]
    for item in checklist:
        add_bullet(doc, item)

    doc.add_page_break()

    # Appendix: Swagger-verified manifest
    add_heading(doc, "Appendix A — Swagger-Verified Endpoint Manifest (74 routes)", 1)
    add_para(doc, "Source: GET /swagger/v1/swagger.json — final authority. Cross-checked with Controllers.")
    swagger_routes = [
        ("GET", "/api/Account/Me"),
        ("PUT", "/api/Account/Me"),
        ("PUT", "/api/Account/update-username"),
        ("POST", "/api/Account/register"),
        ("POST", "/api/Account/LogIn"),
        ("POST", "/api/Account/LogOut"),
        ("POST", "/api/Account/refresh-token"),
        ("POST", "/api/Account/create-admin"),
        ("POST", "/api/Account/sign-in-google"),
        ("POST", "/api/Account/forgot-password/request-code"),
        ("POST", "/api/Account/forgot-password/complete"),
        ("POST", "/api/Account/change-password"),
        ("POST", "/api/Account/set-password"),
        ("POST", "/api/Account/mark-first-login-done"),
        ("POST", "/api/File/Upload"),
        ("GET", "/api/categories"),
        ("GET", "/api/categories/tree"),
        ("GET", "/api/categories/{id}"),
        ("POST", "/api/admin/categories"),
        ("PUT", "/api/admin/categories/{id}"),
        ("DELETE", "/api/admin/categories/{id}"),
        ("GET", "/api/products"),
        ("GET", "/api/products/featured"),
        ("GET", "/api/products/new-arrivals"),
        ("GET", "/api/products/search"),
        ("GET", "/api/products/{id}"),
        ("POST", "/api/admin/products"),
        ("PUT", "/api/admin/products/{id}"),
        ("DELETE", "/api/admin/products/{id}"),
        ("GET", "/api/products/{productId}/images"),
        ("POST", "/api/admin/products/{productId}/images"),
        ("POST", "/api/admin/products/{productId}/images/variant/{variantId}"),
        ("PUT", "/api/admin/products/{productId}/images/{imageId}/primary"),
        ("PUT", "/api/admin/products/{productId}/images/reorder"),
        ("DELETE", "/api/admin/products/{productId}/images/{imageId}"),
        ("GET", "/api/inventory/{variantId}/status"),
        ("GET", "/api/admin/inventory"),
        ("GET", "/api/admin/inventory/low-stock"),
        ("GET", "/api/admin/inventory/{variantId}"),
        ("GET", "/api/admin/inventory/{variantId}/movements"),
        ("POST", "/api/admin/inventory/{variantId}/adjust"),
        ("POST", "/api/admin/inventory/{variantId}/reserve"),
        ("POST", "/api/admin/inventory/{variantId}/release"),
        ("GET", "/api/cart"),
        ("POST", "/api/cart/items"),
        ("PUT", "/api/cart/items/{variantId}"),
        ("DELETE", "/api/cart/items/{variantId}"),
        ("DELETE", "/api/cart"),
        ("POST", "/api/cart/merge"),
        ("POST", "/api/checkout"),
        ("GET", "/api/orders"),
        ("GET", "/api/orders/{id}"),
        ("POST", "/api/orders/{id}/cancel"),
        ("GET", "/api/admin/orders"),
        ("GET", "/api/admin/orders/{id}"),
        ("POST", "/api/admin/orders/{id}/processing"),
        ("POST", "/api/admin/orders/{id}/shipped"),
        ("POST", "/api/admin/orders/{id}/delivered"),
        ("POST", "/api/admin/orders/{id}/cancel"),
        ("POST", "/api/admin/orders/{id}/refund"),
        ("POST", "/api/payments/{orderId}/proof"),
        ("GET", "/api/payments/{orderId}"),
        ("GET", "/api/admin/payments/pending"),
        ("GET", "/api/admin/payments/{paymentId}"),
        ("POST", "/api/admin/payments/{paymentId}/approve"),
        ("POST", "/api/admin/payments/{paymentId}/reject"),
        ("GET", "/api/recommendations/trending"),
        ("GET", "/api/recommendations/recently-viewed"),
        ("GET", "/api/recommendations/similar/{productId}"),
        ("GET", "/api/recommendations/frequently-bought-together/{productId}"),
        ("GET", "/api/Notifications"),
        ("GET", "/api/Notifications/unread-count"),
        ("POST", "/api/Notifications/read/{id}"),
        ("POST", "/api/Notifications/read-all"),
        ("GET", "/api/SystemSettings"),
        ("PUT", "/api/SystemSettings"),
    ]
    add_table(doc, ["#", "Method", "Path"], [
        (str(i + 1), method, path) for i, (method, path) in enumerate(swagger_routes)
    ])
    add_para(doc, "SignalR (not in Swagger paths): /hubs/notifications — events ReceiveNotification, UnreadCountUpdated")

    doc.add_page_break()
    add_heading(doc, "Document Information", 1)
    add_para(doc, "Cairo Bags — Frontend API Integration Architecture")
    add_para(doc, "Integration architecture only. No React components or pages included.")
    add_para(doc, f"Output file: {OUTPUT}")
    add_para(doc, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    return doc


if __name__ == "__main__":
    print("Building Cairo Bags API Integration Architecture report...")
    document = build_document()
    document.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Size: {os.path.getsize(OUTPUT):,} bytes")
