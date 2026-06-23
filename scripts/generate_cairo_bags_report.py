# -*- coding: utf-8 -*-
"""Generate Cairo Bags architecture Word report on Desktop."""
import io
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")
OUTPUT = os.path.join(DESKTOP, "Cairo_Bags_Architecture_Report.docx")


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def add_table(doc, headers, rows, header_color="1A365D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_code_block(doc, text, title=None):
    if title:
        add_para(doc, title, bold=True, size=10)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # light gray background via shading on paragraph - skip for simplicity
    doc.add_paragraph()


def _draw_box(ax, x, y, w, h, text, color="#E8F0FE", edge="#1A365D", fs=8):
    from matplotlib.patches import FancyBboxPatch
    box = FancyBboxPatch((x - w/2, y - h/2), w, h, boxstyle="round,pad=0.02",
                         facecolor=color, edgecolor=edge, linewidth=1.2)
    ax.add_patch(box)
    ax.text(x, y, text, ha="center", va="center", fontsize=fs, fontweight="bold",
            color="#1A365D", wrap=True)


def _arrow(ax, x1, y1, x2, y2):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="->", color="#4A5568", lw=1.2))


def make_current_arch_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 6); ax.axis("off")
    _draw_box(ax, 5, 5.3, 2.5, 0.6, "React Frontend", "#FFF3E0")
    _draw_box(ax, 5, 3.8, 4.0, 1.2, "CommunityHelp API\n(net10.0)", "#E8F0FE")
    _draw_box(ax, 3.2, 3.5, 1.4, 0.5, "Controllers", "#FFFFFF", fs=7)
    _draw_box(ax, 5.0, 3.5, 1.2, 0.5, "Hubs", "#FFFFFF", fs=7)
    _draw_box(ax, 6.8, 3.5, 1.4, 0.5, "Services", "#FFFFFF", fs=7)
    _draw_box(ax, 5.0, 3.0, 2.0, 0.4, "CommunityHelperContext", "#D4EDDA", fs=7)
    _draw_box(ax, 1.5, 1.2, 1.8, 0.6, "SQL Server", "#F8D7DA")
    _draw_box(ax, 4.0, 1.2, 1.6, 0.6, "Google OAuth", "#F8D7DA")
    _draw_box(ax, 6.0, 1.2, 1.6, 0.6, "AI KYC", "#F8D7DA")
    _draw_box(ax, 8.0, 1.2, 1.6, 0.6, "SMTP Email", "#F8D7DA")
    _arrow(ax, 5, 5.0, 5, 4.5)
    _arrow(ax, 5, 3.2, 5, 2.8)
    _arrow(ax, 4.2, 2.8, 1.5, 1.5)
    _arrow(ax, 4.8, 2.8, 4.0, 1.5)
    _arrow(ax, 5.2, 2.8, 6.0, 1.5)
    _arrow(ax, 5.8, 2.8, 8.0, 1.5)
    ax.set_title("Current Community Help Architecture", fontsize=12, fontweight="bold", color="#1A365D")
    buf = io.BytesIO()
    plt.tight_layout(); plt.savefig(buf, format="png", dpi=150, bbox_inches="tight"); plt.close()
    buf.seek(0)
    return buf.read()


def make_target_modules_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")
    groups = [
        (1.5, 3.5, "Catalog", ["Products", "Categories", "Variants", "Inventory"], "#E8F0FE"),
        (5.0, 3.5, "Commerce", ["Cart", "Checkout", "Orders", "Payments"], "#D4EDDA"),
        (8.5, 3.5, "Customer", ["Addresses", "Wishlist", "Reviews"], "#FFF3E0"),
        (3.0, 1.2, "Fulfillment", ["Shipping", "Returns"], "#F3E8FF"),
        (7.0, 1.2, "Intelligence", ["Recommendations", "Analytics"], "#E0F7FA"),
    ]
    for x, y, title, items, color in groups:
        _draw_box(ax, x, y + 0.5, 2.2, 0.5, title, color, fs=9)
        for i, item in enumerate(items):
            _draw_box(ax, x, y - 0.3 - i*0.55, 1.8, 0.4, item, "#FFFFFF", fs=7)
    _arrow(ax, 2.5, 2.5, 4.0, 2.0)
    _arrow(ax, 5.5, 2.5, 4.5, 2.0)
    ax.set_title("Cairo Bags Target Modules", fontsize=12, fontweight="bold", color="#1A365D")
    buf = io.BytesIO()
    plt.tight_layout(); plt.savefig(buf, format="png", dpi=150, bbox_inches="tight"); plt.close()
    buf.seek(0)
    return buf.read()


def make_erd_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 11); ax.set_ylim(0, 7); ax.axis("off")
    entities = [
        (1.2, 6.0, "AspNetUsers"), (3.5, 6.0, "CustomerProfiles"),
        (1.2, 4.8, "ShippingAddresses"), (5.5, 6.0, "Categories"),
        (7.5, 6.0, "Products"), (9.5, 6.0, "ProductVariants"),
        (7.5, 4.8, "ProductTranslations"), (9.5, 4.8, "InventoryItems"),
        (2.0, 3.2, "Carts"), (4.0, 3.2, "CartItems"),
        (6.0, 3.2, "Orders"), (8.0, 3.2, "OrderItems"),
        (6.0, 2.0, "OrderPayments"), (8.0, 2.0, "PaymentProofImages"),
        (4.0, 2.0, "Wishlists"), (2.0, 2.0, "ProductReviews"),
        (6.0, 0.8, "Shipments"), (9.5, 3.2, "Coupons"),
    ]
    for x, y, name in entities:
        _draw_box(ax, x, y, 1.6, 0.45, name, "#E8F0FE", fs=6.5)
    pairs = [
        ((1.2,5.75),(3.5,5.75)), ((1.2,5.55),(1.2,5.05)), ((5.5,5.75),(7.5,5.75)),
        ((7.5,5.75),(9.5,5.75)), ((9.5,5.75),(9.5,5.05)), ((7.5,5.55),(7.5,5.05)),
        ((2.0,2.95),(4.0,2.95)), ((6.0,2.95),(8.0,2.95)), ((6.0,2.75),(6.0,2.25)),
        ((8.0,2.75),(8.0,2.25)), ((6.0,1.75),(6.0,1.05)),
    ]
    for (x1,y1),(x2,y2) in pairs:
        _arrow(ax, x1, y1, x2, y2)
    ax.set_title("Main ERD — Core Entities", fontsize=12, fontweight="bold", color="#1A365D")
    buf = io.BytesIO()
    plt.tight_layout(); plt.savefig(buf, format="png", dpi=150, bbox_inches="tight"); plt.close()
    buf.seek(0)
    return buf.read()


def make_domain_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5.5); ax.axis("off")
    domains = [
        (2.0, 4.5, "Identity & Customer", ["Users", "Profiles", "Addresses"], "#E8F0FE"),
        (5.0, 4.5, "Catalog", ["Categories", "Products", "Variants", "Images"], "#D4EDDA"),
        (8.0, 4.5, "Commerce", ["Carts", "Orders", "Payments"], "#FFF3E0"),
        (3.5, 1.8, "Operations", ["Inventory", "Shipments"], "#F3E8FF"),
        (6.5, 1.8, "Marketing", ["Coupons", "Promotions", "Banners"], "#FCE4EC"),
        (5.0, 0.5, "Intelligence", ["Recommendations", "Analytics"], "#E0F7FA"),
    ]
    for x, y, title, items, color in domains:
        _draw_box(ax, x, y, 2.4, 0.5, title, color, fs=8)
        for i, item in enumerate(items):
            _draw_box(ax, x, y - 0.6 - i*0.5, 1.8, 0.38, item, "#FFFFFF", fs=7)
    ax.set_title("Domain Grouping — Cairo Bags Database", fontsize=12, fontweight="bold", color="#1A365D")
    buf = io.BytesIO()
    plt.tight_layout(); plt.savefig(buf, format="png", dpi=150, bbox_inches="tight"); plt.close()
    buf.seek(0)
    return buf.read()


def make_inventory_flow_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 3.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 3.5); ax.axis("off")
    steps = ["Customer", "Cart", "StockReservation", "InventoryItems", "Order"]
    xs = [1, 3, 5, 7, 9]
    for x, name in zip(xs, steps):
        _draw_box(ax, x, 1.8, 1.6, 0.55, name, "#E8F0FE", fs=7)
    labels = ["Add to cart", "Reserve stock", "Reserved += qty", "Place order\nOnHand -= qty"]
    for i in range(len(xs)-1):
        _arrow(ax, xs[i]+0.8, 1.8, xs[i+1]-0.8, 1.8)
        ax.text((xs[i]+xs[i+1])/2, 2.5, labels[i], ha="center", fontsize=7, color="#4A5568")
    ax.set_title("Inventory Reservation Flow", fontsize=12, fontweight="bold", color="#1A365D")
    buf = io.BytesIO()
    plt.tight_layout(); plt.savefig(buf, format="png", dpi=150, bbox_inches="tight"); plt.close()
    buf.seek(0)
    return buf.read()


def make_order_status_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(11, 4))
    ax.set_xlim(0, 11); ax.set_ylim(0, 4); ax.axis("off")
    cod_steps = ["Pending", "Processing", "Shipped", "Delivered", "Completed"]
    wallet_steps = ["Pending", "Awaiting\nPayment", "Proof\nSubmitted", "Confirmed", "Processing", "Shipped", "Delivered"]
    for i, s in enumerate(cod_steps):
        _draw_box(ax, 0.8 + i*2.0, 3.0, 1.5, 0.55, s, "#D4EDDA", fs=7)
        if i < len(cod_steps)-1:
            _arrow(ax, 0.8+i*2.0+0.75, 3.0, 0.8+(i+1)*2.0-0.75, 3.0)
    ax.text(5.5, 3.7, "Cash On Delivery Flow", ha="center", fontsize=9, fontweight="bold", color="#276749")
    for i, s in enumerate(wallet_steps):
        _draw_box(ax, 0.5 + i*1.45, 1.2, 1.2, 0.55, s, "#FFF3E0", fs=6)
        if i < len(wallet_steps)-1:
            _arrow(ax, 0.5+i*1.45+0.6, 1.2, 0.5+(i+1)*1.45-0.6, 1.2)
    ax.text(5.5, 1.9, "Wallet Payment Flow (InstaPay, Vodafone Cash, etc.)", ha="center", fontsize=9, fontweight="bold", color="#C05621")
    ax.set_title("Order Status Workflows", fontsize=12, fontweight="bold", color="#1A365D")
    buf = io.BytesIO()
    plt.tight_layout(); plt.savefig(buf, format="png", dpi=150, bbox_inches="tight"); plt.close()
    buf.seek(0)
    return buf.read()


DIAGRAM_GENERATORS = {
    "current_arch": make_current_arch_diagram,
    "target_modules": make_target_modules_diagram,
    "erd_main": make_erd_diagram,
    "domain_grouping": make_domain_diagram,
    "inventory_flow": make_inventory_flow_diagram,
    "order_status": make_order_status_diagram,
}


def add_diagram_image(doc, title, generator_key, mermaid_source=None, description=None):
    add_para(doc, title, bold=True, size=12)
    if description:
        add_para(doc, description, italic=True, size=10)
    gen = DIAGRAM_GENERATORS.get(generator_key)
    if gen:
        png = gen()
        stream = io.BytesIO(png)
        doc.add_picture(stream, width=Inches(6.2))
        doc.add_paragraph()
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Figure: {title}")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if mermaid_source:
        add_code_block(doc, mermaid_source, "Diagram source (Mermaid reference):")
    doc.add_paragraph()


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── TITLE PAGE ──
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Cairo Bags")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Complete Architecture & Database Design Report")
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = s2.add_run("Community Help Platform → Premium E-Commerce Migration")
    r3.font.size = Pt(14)
    r3.italic = True

    s3 = doc.add_paragraph()
    s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = s3.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    r4.font.size = Pt(11)

    doc.add_page_break()

    # ── TABLE OF CONTENTS placeholder ──
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "PART 1 — Migration Analysis Report",
        "  1. Current Architecture Overview",
        "  2. Existing Features",
        "  3. Reusable Components",
        "  4. Components That Need Modification",
        "  5. Components To Remove",
        "  6. New Modules Required",
        "  7. Risks And Dependencies",
        "PART 2 — Complete Database Architecture",
        "  8. Design Principles & Enums",
        "  9. Entity Relationship Diagrams",
        "  10. All Table Definitions",
        "  11. Relationships & Indexes",
        "  12. Payment & Order Workflows",
        "  13. Analytics & Recommendations",
        "  14. Seed Data & Migration Mapping",
    ]
    for item in toc_items:
        add_para(doc, item, size=11)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PART 1 — MIGRATION ANALYSIS
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "PART 1 — Migration Analysis Report", 0)
    add_para(doc,
        "This section analyzes the existing Community Help Platform backend and outlines "
        "the migration path to Cairo Bags, a premium e-commerce platform for bags.")

    # 1. Architecture Overview
    add_heading(doc, "1. Current Architecture Overview", 1)

    add_heading(doc, "1.1 Solution Structure", 2)
    add_table(doc,
        ["Item", "Details"],
        [
            ["Workspace root", r"d:\Cairo Bags System"],
            ["Backend root", r"d:\Cairo Bags System\CommunityHelp"],
            ["Primary solution", "CommunityHelp.sln"],
            ["Active project", "CommunityHelp.csproj (only project)"],
            ["Target framework", "net10.0"],
            ["Database", "SQL Server (CommunityHelpDb)"],
            ["Composition root", "Program.cs"],
        ])

    add_heading(doc, "1.2 Folder Structure", 2)
    add_code_block(doc, """CommunityHelp/
├── Program.cs              # DI, middleware, migrations
├── Controllers/            # 12 active REST controllers
├── Models/                 # EF entities + DbContext
├── Service/                # Business services
├── Dto/                    # 58 request/response DTOs
├── Mapper/                 # 10 manual static mappers
├── Hubs/                   # 5 SignalR hubs
├── Helpers/                # 4 utility classes
├── Migrations/             # 44+ EF Core migrations
└── appsettings.json""")

    add_heading(doc, "1.3 Logical Layers", 2)
    add_table(doc,
        ["Layer", "Location", "Responsibility"],
        [
            ["API / Presentation", "Controllers/, Hubs/", "HTTP endpoints, SignalR"],
            ["Application", "Service/", "Token, email, notifications, AI verification"],
            ["Domain + Persistence", "Models/", "Entities, CommunityHelperContext"],
            ["Contracts", "Dto/", "API input/output models"],
            ["Mapping", "Mapper/", "Manual entity ↔ DTO conversion"],
            ["Cross-cutting", "Helpers/, Program.cs", "Auth, CORS, rate limiting"],
        ])

    add_heading(doc, "1.4 Patterns Used", 2)
    add_table(doc,
        ["Pattern", "Status", "Notes"],
        [
            ["Monolithic Web API", "✅", "Single deployable unit"],
            ["DTO pattern", "✅", "58 DTO files"],
            ["Dependency Injection", "✅", "Program.cs registrations"],
            ["ASP.NET Core Identity", "✅", "User : IdentityUser"],
            ["JWT + Refresh tokens", "✅", "TokenService"],
            ["Role-based authorization", "✅", "[Authorize(Roles)]"],
            ["Google OAuth", "✅", "GoogleSignInService"],
            ["EF Core Code First", "✅", "SQL Server"],
            ["SignalR", "✅", "Chat, events, notifications"],
            ["CQRS / MediatR", "❌", "Not used"],
            ["Repository pattern", "❌", "Direct DbContext"],
            ["Clean Architecture", "❌", "Single project"],
        ])

    add_heading(doc, "1.5 NuGet Dependencies", 2)
    add_table(doc,
        ["Package", "Version", "Purpose"],
        [
            ["Microsoft.AspNetCore.Authentication.JwtBearer", "10.0.5", "JWT auth"],
            ["Microsoft.AspNetCore.Identity.EntityFrameworkCore", "8.0.0", "Identity"],
            ["Microsoft.EntityFrameworkCore.SqlServer", "9.0.4", "ORM"],
            ["Google.Apis.Auth", "1.69.0", "Google OAuth"],
            ["Swashbuckle.AspNetCore", "6.6.2", "Swagger"],
            ["Newtonsoft.Json", "13.0.4", "JSON"],
        ])

    add_heading(doc, "1.6 Current Architecture Diagram", 2)
    current_arch = """flowchart TB
    subgraph Client
        React[React Frontend]
    end
    subgraph API[CommunityHelp net10.0]
        Controllers[Controllers]
        Hubs[SignalR Hubs]
        Services[Services]
        DbContext[CommunityHelperContext]
    end
    subgraph External
        SQL[(SQL Server)]
        Google[Google OAuth]
        AI[AI KYC Service]
        SMTP[Email SMTP]
    end
    React -->|JWT REST| Controllers
    React -->|WebSocket| Hubs
    Controllers --> Services
    Controllers --> DbContext
    Hubs --> DbContext
    Services --> DbContext
    DbContext --> SQL"""
    add_diagram_image(doc, "Current Community Help Architecture", "current_arch",
                      current_arch, "Monolithic ASP.NET Core API with direct EF Core access.")

    # 2. Existing Features
    add_heading(doc, "2. Existing Features (Implemented Modules)", 1)

    modules = [
        ("Authentication & Account", "AccountController", "Register, login, Google OAuth, JWT refresh, password reset, profile"),
        ("Identity Verification (KYC)", "VerificationController", "AI + admin document verification, fraud detection"),
        ("NGO Staff Management", "NGOStaffController", "NGO Manager/Employee hierarchy"),
        ("Community Events", "EventController, VolunteerController", "Event CRUD, volunteering, SignalR updates"),
        ("Blood Donation", "BloodDonationController", "Donation request lifecycle, donor reputation"),
        ("Issue Reporting", "IssueController", "Citizen issues, admin moderation"),
        ("Messaging", "MessagesController, MessageHub", "User-to-user real-time chat"),
        ("Notifications", "NotificationsController, NotificationHub", "In-app notifications with deduplication"),
        ("User Profiles", "UsersController", "Stats, XP leaderboard, admin details"),
        ("System Settings", "SystemSettingsController", "Maintenance mode, beta features"),
        ("File Upload", "FileController", "Image/audio upload to local storage"),
        ("Legacy Animal2", "DashBoardController (disabled)", "Orphaned DTOs, mappers, broken solution"),
    ]
    add_table(doc, ["Module", "Controller/Hub", "Description"], modules)

    add_heading(doc, "2.1 Database Entities (Current)", 2)
    add_table(doc,
        ["Entity", "Domain"],
        [
            ["User", "Identity + KYC + NGO hierarchy"],
            ["UserCategory", "Admin, Citizen, NGO Manager, NGO Employee"],
            ["Event, EventVolunteer", "Community events"],
            ["Issue", "Community issues"],
            ["BloodDonationRequest, BloodDonation", "Blood donation"],
            ["Message", "Chat"],
            ["Notification", "In-app alerts"],
            ["SystemSetting", "Platform config"],
            ["PasswordResetOtp", "OTP reset"],
            ["VerificationAttempt", "KYC audit"],
        ])

    # 3. Reusable
    add_heading(doc, "3. Reusable Components", 1)
    add_table(doc,
        ["Component", "Reuse %", "Notes"],
        [
            ["TokenService + JWT config", "~85%", "JWT + refresh tokens"],
            ["GoogleSignInService", "~85%", "Social login for customers"],
            ["PasswordResetService + OTP", "~85%", "Password recovery"],
            ["EmailService (IEmailService)", "~70%", "Templates need rewrite"],
            ["FileController upload pattern", "~60%", "Product images"],
            ["NotificationService + Hub", "~50%", "Types/targets change"],
            ["CORS, rate limiting, Swagger", "~90%", "Config updates only"],
            ["NameIdentifierUserIdProvider", "~90%", "SignalR user mapping"],
            ["Domain entities & controllers", "~10%", "Mostly replaced"],
        ])

    # 4. Modification
    add_heading(doc, "4. Components That Need Modification", 1)
    mods = [
        ("User model", "Remove NGO/KYC fields; add addresses, loyalty, wishlist"),
        ("Roles", "Admin, Citizen → Admin, Customer; remove NGO roles"),
        ("AccountController", "RegisterCustomer, remove verification/NGO flows"),
        ("Notifications", "New types: Order, Shipment, Promotion, Review"),
        ("UsersController", "Order stats instead of XP/donations"),
        ("Messaging", "Repurpose as customer support chat"),
        ("SystemSettings", "Add tax, shipping threshold, store name"),
        ("DbContext", "Rename, new DbSets, greenfield migrations"),
        ("appsettings.json", "Remove AiVerification; add Payments, Shipping"),
    ]
    for name, desc in mods:
        add_bullet(doc, f"{name}: {desc}")

    # 5. Remove
    add_heading(doc, "5. Components To Remove Completely", 1)
    add_table(doc,
        ["Category", "Items to Delete"],
        [
            ["Controllers", "BloodDonation, Event, Volunteer, Issue, NGOStaff, Verification, DashBoard"],
            ["Hubs", "BloodDonationHub, EventsHub, IssueHub"],
            ["Services", "AiVerification, EventStatusService, FraudDetection, KYC helpers"],
            ["Entities", "Event, EventVolunteer, Issue, BloodDonation*, VerificationAttempt, UserCategory"],
            ["DTOs", "~35 files (Blood, Event, Issue, Volunteer, Legacy Animal2, KYC)"],
            ["Mappers", "8 files (all community + legacy mappers)"],
            ["Config", "AiVerification section"],
            ["Solutions", "Animal2.sln, aaa.sln"],
        ])

    # 6. New Modules
    add_heading(doc, "6. New Modules Required for Cairo Bags", 1)
    add_table(doc,
        ["Module", "Priority", "Description"],
        [
            ["Product Catalog", "P0", "Products, categories, variants, images, collections"],
            ["Inventory Management", "P0", "Stock levels, movements, reservations"],
            ["Shopping Cart", "P0", "User/guest cart with merge"],
            ["Checkout & Orders", "P0", "Order workflow, snapshots"],
            ["Payments (Egyptian)", "P0", "COD + wallet methods with proof upload"],
            ["Shipping", "P0", "Cairo/Giza/Other zones, rates, tracking"],
            ["Wishlist", "P1", "Saved products"],
            ["Reviews & Ratings", "P1", "Verified purchase reviews"],
            ["Coupons & Promotions", "P1", "Discount codes, automatic promos"],
            ["Banners", "P1", "Homepage/promotional banners"],
            ["Notifications (adapted)", "P1", "Order/shipment/promo alerts"],
            ["Recommendation Engine", "P2", "Trending, cross-sell, personalized"],
            ["Analytics Dashboard", "P2", "Sales, products, payments, zones"],
            ["Returns & Refunds", "P2", "Return workflow"],
            ["Admin Dashboard", "P1", "Catalog, orders, payments review"],
        ])

    new_modules_diagram = """flowchart LR
    subgraph Catalog
        P[Products]
        C[Categories]
        V[Variants]
        I[Inventory]
    end
    subgraph Commerce
        Cart[Cart]
        CO[Checkout]
        O[Orders]
        Pay[Payments]
    end
    subgraph Customer
        A[Addresses]
        W[Wishlist]
        R[Reviews]
    end
    P --> Cart
    V --> Cart
    I --> Cart
    Cart --> CO --> O --> Pay"""
    add_diagram_image(doc, "Cairo Bags Target Module Map", "target_modules", new_modules_diagram)

    # 7. Risks
    add_heading(doc, "7. Risks And Dependencies", 1)
    add_table(doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["Full domain pivot", "70%+ schema invalid", "Greenfield database"],
            ["No payment module", "Cannot transact", "Priority P0 development"],
            ["No inventory concurrency", "Overselling", "Stock reservations + RowVersion"],
            ["44+ legacy migrations", "Schema evolution blocked", "New InitialCreate migration"],
            ["Monolithic coupling", "Hard to refactor", "Introduce service layer"],
            ["Package version mismatch", "Runtime issues", "Align net10/EF/Identity versions"],
            ["KYC PII in database", "Privacy risk", "Secure purge on migration"],
            ["SignalR client coupling", "Frontend breaks", "API versioning"],
            ["No test project", "No regression safety", "Add tests before refactor"],
        ])

    add_heading(doc, "7.1 Migration Phases", 2)
    add_table(doc,
        ["Phase", "Focus", "Duration"],
        [
            ["Phase 0", "Rename solution, align packages, new DbContext", "1–2 weeks"],
            ["Phase 1", "Strip community modules, KYC, legacy", "1 week"],
            ["Phase 2", "Auth & Customer model refactor", "1–2 weeks"],
            ["Phase 3", "Product catalog & inventory", "2–3 weeks"],
            ["Phase 4", "Cart, checkout, orders, payments", "3–4 weeks"],
            ["Phase 5", "Shipping, returns, notifications", "2 weeks"],
            ["Phase 6", "Reviews, wishlist, loyalty, search, admin", "2–3 weeks"],
            ["Phase 7", "Security audit, PII purge, cloud storage", "1–2 weeks"],
        ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # PART 2 — DATABASE ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    add_heading(doc, "PART 2 — Complete Database Architecture", 0)
    add_para(doc,
        "Production-ready database design for Cairo Bags premium e-commerce platform. "
        "SQL Server, bilingual (Arabic/English), Egyptian payment methods and shipping zones.")

    # 8. Design Principles
    add_heading(doc, "8. Design Principles & Enums", 1)
    add_table(doc,
        ["Principle", "Decision"],
        [
            ["Primary keys", "INT IDENTITY for business; NVARCHAR(450) for users"],
            ["Money", "DECIMAL(18,2)"],
            ["Timestamps", "DATETIME2(7) UTC"],
            ["Localization", "Dedicated *Translations tables"],
            ["Audit fields", "CreatedAt, UpdatedAt, CreatedBy, UpdatedBy"],
            ["Concurrency", "ROWVERSION on InventoryItems"],
            ["Soft delete", "IsDeleted on addresses, products"],
        ])

    add_heading(doc, "8.1 Core Enums", 2)
    add_table(doc,
        ["Enum", "Key Values"],
        [
            ["ProductStatus", "Draft, Active, Archived"],
            ["OrderStatus", "Pending → AwaitingPayment → ... → Completed"],
            ["PaymentMethodType", "COD, InstaPay, Vodafone Cash, Orange Cash, Etisalat Cash, WE Pay"],
            ["PaymentStatus", "Pending, ProofSubmitted, UnderReview, Confirmed, Rejected, Refunded"],
            ["ShippingZoneType", "Cairo, Giza, OtherGovernorates"],
            ["CouponType", "Percentage, FixedAmount, FreeShipping"],
            ["NotificationType", "Order, Payment, Shipment, Promotion, Review, System"],
            ["ReviewStatus", "Pending, Approved, Rejected"],
            ["InventoryMovementType", "Sale, Return, Adjustment, Reservation, ReservationRelease"],
        ])

    add_heading(doc, "8.2 Payment Methods", 2)
    add_table(doc,
        ["Id", "Code", "Method", "Requires Proof", "Admin Confirm"],
        [
            ["1", "COD", "Cash On Delivery", "No", "No"],
            ["2", "INSTAPAY", "InstaPay", "Yes", "Yes"],
            ["3", "VODAFONE_CASH", "Vodafone Cash", "Yes", "Yes"],
            ["4", "ORANGE_CASH", "Orange Cash", "Yes", "Yes"],
            ["5", "ETISALAT_CASH", "Etisalat Cash", "Yes", "Yes"],
            ["6", "WE_PAY", "WE Pay", "Yes", "Yes"],
        ])

    add_heading(doc, "8.3 Shipping Zones", 2)
    add_table(doc,
        ["Zone", "Code", "Governorates"],
        [
            ["Cairo", "CAIRO", "Cairo"],
            ["Giza", "GIZA", "Giza"],
            ["Other Governorates", "OTHER", "All remaining 25 governorates"],
        ])

    # 9. ERD
    add_heading(doc, "9. Entity Relationship Diagrams", 1)

    erd_main = """erDiagram
    AspNetUsers ||--o| CustomerProfiles : has
    AspNetUsers ||--o{ ShippingAddresses : owns
    AspNetUsers ||--o{ Orders : places
    Categories ||--o{ Products : contains
    Products ||--o{ ProductVariants : has
    Products ||--o{ ProductTranslations : localized
    ProductVariants ||--o{ InventoryItems : tracked
    Carts ||--o{ CartItems : contains
    Orders ||--o{ OrderItems : contains
    Orders ||--o{ OrderPayments : has
    OrderPayments ||--o{ PaymentProofImages : proofs
    Orders ||--o| Shipments : fulfilled
    Coupons ||--o{ CouponUsages : used
    Products ||--o{ ProductReviews : reviewed
    Products ||--o{ ProductRecommendations : recommends"""
    add_diagram_image(doc, "Main ERD — Core Entities", "erd_main", erd_main)

    erd_domain = """flowchart TB
    subgraph Identity[Identity and Customer]
        U[AspNetUsers]
        CP[CustomerProfiles]
        SA[ShippingAddresses]
    end
    subgraph Catalog[Catalog]
        CAT[Categories]
        P[Products]
        PV[ProductVariants]
    end
    subgraph Commerce[Commerce]
        CART[Carts]
        O[Orders]
        OP[OrderPayments]
    end
    subgraph Ops[Operations]
        INV[InventoryItems]
        SH[Shipments]
    end
    U --> CP
    U --> SA
    U --> CART
    U --> O
    CAT --> P --> PV --> INV
    CART --> O --> OP
    O --> SH"""
    add_diagram_image(doc, "Domain Grouping Diagram", "domain_grouping", erd_domain)

  # 10. Table Definitions - grouped
    add_heading(doc, "10. All Table Definitions", 1)
    add_para(doc, "Total: 81 tables (including 7 ASP.NET Identity tables). All business tables include audit fields: CreatedAt, UpdatedAt, CreatedBy, UpdatedBy.", bold=True)

    table_groups = [
        ("Identity & Access (8 tables)", [
            "AspNetUsers — Core auth; extends with PreferredLanguage, RefreshToken, IsActive",
            "AspNetRoles — Admin, Customer",
            "AspNetUserRoles, AspNetUserClaims, AspNetRoleClaims, AspNetUserLogins, AspNetUserTokens — Standard Identity",
            "CustomerProfiles — 1:1 user extension: name, loyalty, total spent",
            "AdminProfiles — Admin metadata",
            "PasswordResetOtps — Hashed OTP for password reset",
        ]),
        ("Localization & Reference (3 tables)", [
            "Languages — en, ar with RTL flag",
            "Currencies — EGP default",
            "Governorates — 27 Egyptian governorates mapped to shipping zones",
        ]),
        ("Shipping (5 tables)", [
            "ShippingZones — Cairo, Giza, OtherGovernorates",
            "ShippingZoneTranslations — Bilingual zone names",
            "ShippingRates — Base fee, free shipping threshold per zone",
            "ShippingAddresses — Customer address book with governorate, district, landmark",
        ]),
        ("Catalog (14 tables)", [
            "Categories + CategoryTranslations — Hierarchical taxonomy, bilingual SEO slugs",
            "Products — Base price, status, ratings, material, featured flags",
            "ProductTranslations — Name, description, slug (ar/en)",
            "ProductImages — Gallery with variant-specific images",
            "ProductCategoryMappings — Many-to-many categories",
            "AttributeDefinitions + Translations + Options — Color, Size, Material",
            "ProductVariants — SKU-level sellable unit with price",
            "VariantAttributeValues — Variant attribute combinations",
            "Collections + CollectionTranslations + CollectionProducts — Curated lookbooks",
        ]),
        ("Inventory (4 tables)", [
            "Warehouses — Stock locations (MAIN default)",
            "InventoryItems — OnHand, Reserved, Available per variant/warehouse + RowVersion",
            "InventoryMovements — Full audit trail of stock changes",
            "StockReservations — Checkout holds with expiry",
        ]),
        ("Cart & Wishlist (4 tables)", [
            "Carts — User or guest session, coupon, subtotal",
            "CartItems — Variant + quantity + unit price",
            "Wishlists — One per user",
            "WishlistItems — Saved product/variant",
        ]),
        ("Orders (4 tables)", [
            "Orders — OrderNumber, status, totals, payment method, shipping zone",
            "OrderItems — Snapshot: SKU, names (ar/en), price, image at purchase time",
            "OrderAddresses — Frozen delivery address snapshot",
            "OrderStatusHistories — Full status transition audit",
        ]),
        ("Payments (4 tables)", [
            "PaymentMethods + PaymentMethodTranslations — Egyptian payment options",
            "OrderPayments — Payment status, sender wallet, transaction ref, admin review",
            "PaymentProofImages — Transfer screenshot uploads",
        ]),
        ("Shipments (2 tables)", [
            "Shipments — Carrier, tracking, delivery dates",
            "ShipmentTrackingEvents — Customer-visible timeline",
        ]),
        ("Coupons (3 tables)", [
            "Coupons — Code, type, scope, limits, validity",
            "CouponScopes — Category/product restrictions",
            "CouponUsages — Per-user usage tracking",
        ]),
        ("Promotions (5 tables)", [
            "Promotions + PromotionTranslations — Automatic discounts",
            "PromotionProducts, PromotionCategories — Scoped targets",
            "PromotionRules — BuyXGetY, min quantity rules (JSON config)",
        ]),
        ("Reviews (3 tables)", [
            "ProductReviews — Rating, comment, verified purchase, moderation",
            "ReviewImages — Customer photos",
            "ReviewHelpfulVotes — Helpful voting",
        ]),
        ("Banners (3 tables)", [
            "BannerPlacements — HOME_HERO, CATEGORY_TOP, etc.",
            "Banners — Images (ar/en), links, scheduling",
            "BannerTranslations — Title, subtitle, CTA text",
        ]),
        ("Notifications (4 tables)", [
            "NotificationTemplates + Translations — System message templates",
            "Notifications — In-app inbox with deduplication",
            "UserNotificationPreferences — Opt-in per category",
        ]),
        ("Recommendations (5 tables)", [
            "ProductRecommendations — Cross-sell / upsell with score",
            "UserProductViews — Browse behavior",
            "UserProductInteractions — View, cart, wishlist, purchase events",
            "RecommendationRules — Configurable algorithms",
            "TrendingProducts — Precomputed trending lists",
        ]),
        ("Analytics (6 tables)", [
            "DailySalesSummaries — Revenue, orders, AOV per day",
            "ProductPerformanceMetrics — Views, conversion, revenue per product",
            "CategoryPerformanceMetrics — Category-level sales",
            "PaymentMethodAnalytics — Payment method split",
            "ShippingZoneAnalytics — Zone distribution",
            "AnalyticsEvents — Raw event stream",
        ]),
        ("System (2 tables)", [
            "SystemSettings — Store name, tax, shipping threshold, maintenance",
            "AuditLogs — Admin change accountability (JSON old/new values)",
        ]),
    ]

    for group_name, tables in table_groups:
        add_heading(doc, group_name, 2)
        for t in tables:
            add_bullet(doc, t)

    # Key table columns detail
    add_heading(doc, "10.1 Key Table Column Details", 2)

    add_para(doc, "Products", bold=True)
    add_table(doc,
        ["Column", "Type", "Notes"],
        [
            ["Id", "INT PK", "IDENTITY"],
            ["CategoryId", "INT FK", "Primary category"],
            ["SkuPrefix", "NVARCHAR(20)", "e.g. CB-BAG"],
            ["BasePrice", "DECIMAL(18,2)", ""],
            ["Status", "TINYINT", "Draft/Active/Archived"],
            ["AverageRating", "DECIMAL(3,2)", "Denormalized"],
            ["HasVariants", "BIT", ""],
            ["+ audit fields", "", "CreatedAt, UpdatedAt, CreatedBy, UpdatedBy"],
        ])

    add_para(doc, "ProductVariants", bold=True)
    add_table(doc,
        ["Column", "Type", "Notes"],
        [
            ["Id", "INT PK", ""],
            ["ProductId", "INT FK", ""],
            ["Sku", "NVARCHAR(50) UNIQUE", "Full SKU"],
            ["Price", "DECIMAL(18,2)", "Overrides product base"],
            ["Status", "TINYINT", "Active/Inactive"],
            ["IsDefault", "BIT", ""],
        ])

    add_para(doc, "Orders", bold=True)
    add_table(doc,
        ["Column", "Type", "Notes"],
        [
            ["Id", "INT PK", ""],
            ["OrderNumber", "NVARCHAR(20) UNIQUE", "CB-2026-000001"],
            ["UserId", "NVARCHAR(450) FK", ""],
            ["Status", "TINYINT", "OrderStatus workflow"],
            ["Subtotal / Discount / Shipping / Tax / Total", "DECIMAL(18,2)", ""],
            ["PaymentMethodId", "INT FK", ""],
            ["ShippingZoneId", "INT FK", ""],
            ["PlacedAt", "DATETIME2", ""],
        ])

    add_para(doc, "InventoryItems", bold=True)
    add_table(doc,
        ["Column", "Type", "Notes"],
        [
            ["VariantId + WarehouseId", "FK", "UNIQUE composite"],
            ["QuantityOnHand", "INT", "Physical stock"],
            ["QuantityReserved", "INT", "Checkout holds"],
            ["QuantityAvailable", "INT", "OnHand - Reserved"],
            ["RowVersion", "ROWVERSION", "Optimistic concurrency"],
            ["LowStockThreshold", "INT", "Alert trigger"],
        ])

    # 11. Relationships & Indexes
    add_heading(doc, "11. Relationships & Critical Indexes", 1)
    add_table(doc,
        ["Relationship", "Cardinality", "Description"],
        [
            ["User → CustomerProfile", "1:1", "Customer extension"],
            ["User → ShippingAddresses", "1:N", "Address book"],
            ["User → Orders", "1:N", "Purchase history"],
            ["Product → Variants", "1:N", "SKUs"],
            ["Variant → InventoryItems", "1:N", "Per warehouse"],
            ["Cart → CartItems → Variant", "1:N → N:1", "Shopping cart"],
            ["Order → OrderItems", "1:N", "Immutable snapshots"],
            ["Order → OrderPayments", "1:N", "Payment lifecycle"],
            ["OrderPayment → ProofImages", "1:N", "Transfer screenshots"],
            ["Coupon → CouponUsages", "1:N", "Usage tracking"],
            ["Governorate → ShippingZone", "N:1", "Zone mapping"],
        ])

    add_heading(doc, "11.1 Critical Indexes", 2)
    add_table(doc,
        ["Table", "Index", "Purpose"],
        [
            ["ProductTranslations", "UNIQUE (Slug, LanguageCode)", "SEO URLs"],
            ["ProductVariants", "UNIQUE Sku", "SKU lookup"],
            ["InventoryItems", "UNIQUE (VariantId, WarehouseId)", "Stock row"],
            ["Orders", "UNIQUE OrderNumber", "Customer lookup"],
            ["Orders", "(UserId, PlacedAt DESC)", "Order history"],
            ["Coupons", "UNIQUE Code", "Checkout validation"],
            ["Notifications", "(UserId, IsRead, CreatedAt DESC)", "Inbox"],
            ["StockReservations", "(ExpiresAt, Status)", "Expiry cleanup"],
            ["ProductReviews", "(ProductId, Status, Rating)", "Product page"],
        ])

    # 12. Workflows
    add_heading(doc, "12. Payment & Order Workflows", 1)

    add_heading(doc, "12.1 Order Status Workflow", 2)
    add_diagram_image(doc, "Order Status Workflows (COD vs Wallet)", "order_status")
    add_table(doc,
        ["Code", "Status", "Next States"],
        [
            ["1", "Pending", "AwaitingPayment, Processing (COD), Cancelled"],
            ["2", "AwaitingPayment", "PaymentProofSubmitted, Cancelled"],
            ["3", "PaymentProofSubmitted", "PaymentUnderReview"],
            ["4", "PaymentUnderReview", "PaymentConfirmed, AwaitingPayment (reject)"],
            ["5", "PaymentConfirmed", "Processing"],
            ["6", "Processing", "Shipped"],
            ["7", "Shipped", "OutForDelivery"],
            ["8", "OutForDelivery", "Delivered"],
            ["9", "Delivered", "Completed, ReturnRequested"],
            ["10", "Completed", "ReturnRequested"],
            ["11", "Cancelled", "Terminal"],
            ["12-15", "Refund/Return states", "Terminal branches"],
        ])

    add_heading(doc, "12.2 COD vs Wallet Payment Flow", 2)
    add_table(doc,
        ["Step", "Cash On Delivery", "Wallet Methods (InstaPay, etc.)"],
        [
            ["1", "Order → Processing", "Order → AwaitingPayment"],
            ["2", "Ship immediately", "Customer pays externally"],
            ["3", "Collect cash on delivery", "Upload PaymentProofImage"],
            ["4", "Mark Completed", "Admin reviews proof"],
            ["5", "—", "Confirm → Processing → Ship"],
        ])

    inventory_flow = """sequenceDiagram
    participant C as Customer
    participant Cart
    participant SR as StockReservations
    participant II as InventoryItems
    participant O as Orders
    C->>Cart: Add to cart
    Cart->>SR: Create reservation
    SR->>II: QuantityReserved += qty
    C->>O: Place order
    SR->>II: OnHand -= qty, Reserved -= qty"""
    add_diagram_image(doc, "Inventory Reservation Flow", "inventory_flow", inventory_flow)

    # 13. Analytics
    add_heading(doc, "13. Analytics Dashboard & Recommendations", 1)
    add_table(doc,
        ["Dashboard Widget", "Source Tables"],
        [
            ["Revenue today/week/month", "DailySalesSummaries"],
            ["Orders by status", "Orders"],
            ["Top selling products", "ProductPerformanceMetrics"],
            ["Payment method split", "PaymentMethodAnalytics"],
            ["Shipping zone distribution", "ShippingZoneAnalytics"],
            ["Low stock alerts", "InventoryItems"],
            ["Pending payment reviews", "OrderPayments (ProofSubmitted)"],
            ["Review moderation queue", "ProductReviews (Pending)"],
            ["Trending products", "TrendingProducts"],
            ["Personalized recommendations", "ProductRecommendations, UserProductViews"],
        ])

    add_heading(doc, "13.1 Background Jobs", 2)
    add_table(doc,
        ["Job", "Frequency", "Tables"],
        [
            ["Expire stock reservations", "Every 5 min", "StockReservations, InventoryItems"],
            ["Abandon stale carts", "Hourly", "Carts"],
            ["Recompute trending", "Daily", "TrendingProducts"],
            ["Aggregate daily sales", "Daily", "DailySalesSummaries"],
            ["Low stock alerts", "Hourly", "InventoryItems → Notifications"],
            ["Review reminders", "Daily", "Orders → Notifications"],
        ])

    # 14. Seed & Migration
    add_heading(doc, "14. Seed Data & Migration Mapping", 1)

    add_heading(doc, "14.1 Required Seed Data", 2)
    add_table(doc,
        ["Table", "Seed Data"],
        [
            ["Languages", "en, ar"],
            ["Currencies", "EGP"],
            ["AspNetRoles", "Admin, Customer"],
            ["ShippingZones", "Cairo, Giza, Other Governorates"],
            ["Governorates", "27 Egyptian governorates"],
            ["PaymentMethods", "6 Egyptian payment methods"],
            ["Warehouses", "MAIN"],
            ["BannerPlacements", "HOME_HERO, HOME_PROMO, CATEGORY_TOP"],
            ["SystemSettings", "Single Cairo Bags default row"],
        ])

    add_heading(doc, "14.2 Community Help → Cairo Bags Mapping", 2)
    add_table(doc,
        ["Community Help", "Cairo Bags", "Action"],
        [
            ["AspNetUsers / User", "AspNetUsers + CustomerProfiles", "Migrate credentials; strip KYC/NGO"],
            ["PasswordResetOtps", "PasswordResetOtps", "Reuse"],
            ["Notifications", "Notifications", "New schema; truncate old"],
            ["SystemSetting", "SystemSettings", "Extend"],
            ["Events, Blood, Issues, KYC", "—", "Drop entirely"],
            ["—", "All catalog/commerce tables", "New greenfield"],
        ])

    add_heading(doc, "14.3 Complete Table Index (81 tables)", 2)
    all_tables = [
        "AdminProfiles", "AnalyticsEvents", "AspNetRoleClaims", "AspNetRoles",
        "AspNetUserClaims", "AspNetUserLogins", "AspNetUserRoles", "AspNetUsers",
        "AspNetUserTokens", "AttributeDefinitionTranslations", "AttributeDefinitions",
        "AttributeOptionTranslations", "AttributeOptions", "AuditLogs",
        "BannerPlacements", "BannerTranslations", "Banners", "CartItems", "Carts",
        "Categories", "CategoryPerformanceMetrics", "CategoryTranslations",
        "CollectionProducts", "CollectionTranslations", "Collections", "CouponScopes",
        "CouponUsages", "Coupons", "Currencies", "CustomerProfiles",
        "DailySalesSummaries", "Governorates", "InventoryItems", "InventoryMovements",
        "Languages", "NotificationTemplateTranslations", "NotificationTemplates",
        "Notifications", "OrderAddresses", "OrderItems", "OrderPayments",
        "OrderStatusHistories", "Orders", "PasswordResetOtps",
        "PaymentMethodAnalytics", "PaymentMethodTranslations", "PaymentMethods",
        "PaymentProofImages", "ProductCategoryMappings", "ProductImages",
        "ProductPerformanceMetrics", "ProductRecommendations", "ProductReviews",
        "ProductTranslations", "ProductVariants", "Products", "PromotionCategories",
        "PromotionProducts", "PromotionRules", "PromotionTranslations", "Promotions",
        "RecommendationRules", "ReviewHelpfulVotes", "ReviewImages",
        "ShipmentTrackingEvents", "Shipments", "ShippingAddresses", "ShippingRates",
        "ShippingZoneAnalytics", "ShippingZoneTranslations", "ShippingZones",
        "StockReservations", "SystemSettings", "TrendingProducts",
        "UserNotificationPreferences", "UserProductInteractions", "UserProductViews",
        "VariantAttributeValues", "Warehouses", "WishlistItems", "Wishlists",
    ]
    # Print in 4 columns as table rows
    rows = []
    for i in range(0, len(all_tables), 4):
        chunk = all_tables[i:i+4]
        while len(chunk) < 4:
            chunk.append("")
        rows.append(chunk)
    add_table(doc, ["#", "#", "#", "#"], rows)

    # Footer
    doc.add_page_break()
    add_heading(doc, "Document Information", 1)
    add_para(doc, "This document consolidates:")
    add_bullet(doc, "Part 1: Migration Analysis Report (Community Help → Cairo Bags)")
    add_bullet(doc, "Part 2: Complete Database Architecture (81 tables, bilingual, Egyptian payments)")
    add_para(doc, "Status: Architecture documentation only. No code, DbContext, or migrations were generated.")
    add_para(doc, f"File location: {OUTPUT}")

    return doc


if __name__ == "__main__":
    print("Building Cairo Bags Architecture Report...")
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Size: {os.path.getsize(OUTPUT):,} bytes")
