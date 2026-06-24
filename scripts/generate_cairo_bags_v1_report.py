# -*- coding: utf-8 -*-
"""Generate Cairo Bags V1 Database Architecture Word report on Desktop."""
import io
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")
OUTPUT = os.path.join(DESKTOP, "Cairo_Bags_V1_Database_Architecture.docx")


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def add_table(doc, headers, rows, header_color="1A365D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def _draw_box(ax, x, y, w, h, text, color="#E8F0FE", edge="#1A365D", fs=8):
    from matplotlib.patches import FancyBboxPatch
    box = FancyBboxPatch((x - w / 2, y - h / 2), w, h, boxstyle="round,pad=0.02",
                         facecolor=color, edgecolor=edge, linewidth=1.2)
    ax.add_patch(box)
    ax.text(x, y, text, ha="center", va="center", fontsize=fs, fontweight="bold",
            color="#1A365D", wrap=True)


def _arrow(ax, x1, y1, x2, y2):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="->", color="#4A5568", lw=1.2))


def _save_fig(fig):
    buf = io.BytesIO()
    import matplotlib.pyplot as plt
    plt.tight_layout()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def make_v1_module_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.5)
    ax.axis("off")
    groups = [
        (1.8, 4.2, "Identity", ["Auth", "Admin / Customer", "Profiles"], "#E8F0FE"),
        (5.0, 4.2, "Catalog", ["Products", "Variants", "Images", "Categories"], "#D4EDDA"),
        (8.2, 4.2, "Commerce", ["Cart", "Orders", "Payments"], "#FFF3E0"),
        (2.5, 1.8, "Operations", ["Inventory", "Notifications", "Audit Logs"], "#F3E8FF"),
        (5.5, 1.8, "Store", ["Shipping Zones", "Coupons", "Banners"], "#FCE4EC"),
        (8.2, 1.8, "Intelligence", ["Analytics", "Recommendations"], "#E0F7FA"),
    ]
    for x, y, title, items, color in groups:
        _draw_box(ax, x, y + 0.45, 2.0, 0.45, title, color, fs=9)
        for i, item in enumerate(items):
            _draw_box(ax, x, y - 0.2 - i * 0.48, 1.7, 0.36, item, "#FFFFFF", fs=7)
    ax.set_title("Cairo Bags V1 — Module Map (Single Store)", fontsize=12, fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def make_v1_erd_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(11, 7.5))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 7.5)
    ax.axis("off")
    entities = [
        (1.2, 6.5, "Users"), (3.2, 6.5, "CustomerProfiles"),
        (1.2, 5.3, "ShippingAddr"), (5.0, 6.5, "Categories"),
        (7.0, 6.5, "Products"), (9.2, 6.5, "Variants"),
        (7.0, 5.3, "ProdTranslations"), (9.2, 5.3, "Inventory"),
        (2.5, 3.8, "Carts"), (4.5, 3.8, "CartItems"),
        (6.5, 3.8, "Orders"), (8.5, 3.8, "OrderItems"),
        (6.5, 2.6, "OrderPayments"), (8.5, 2.6, "ProofImages"),
        (2.5, 2.6, "Wishlists"), (4.5, 2.6, "Reviews"),
        (1.2, 1.2, "Coupons"), (5.0, 1.2, "ShippingZones"),
        (7.5, 1.2, "Banners"), (9.2, 1.2, "Trending"),
    ]
    for x, y, name in entities:
        _draw_box(ax, x, y, 1.5, 0.42, name, "#E8F0FE", fs=6.5)
    pairs = [
        ((1.2, 6.28), (3.2, 6.28)), ((5.0, 6.28), (7.0, 6.28)), ((7.0, 6.28), (9.2, 6.28)),
        ((9.2, 6.28), (9.2, 5.52)), ((7.0, 5.08), (7.0, 5.52)),
        ((2.5, 3.58), (4.5, 3.58)), ((6.5, 3.58), (8.5, 3.58)),
        ((6.5, 3.38), (6.5, 2.82)), ((8.5, 3.38), (8.5, 2.82)),
    ]
    for a, b in pairs:
        _arrow(ax, *a, *b)
    ax.set_title("V1 ERD — Core Entities (47 Tables)", fontsize=12, fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def make_v1_domain_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.5)
    ax.axis("off")
    domains = [
        (2.0, 3.5, "Catalog", ["Products + Translations", "Variants (Color)", "Images"], "#D4EDDA"),
        (5.0, 3.5, "Commerce", ["Cart", "Orders", "Payments + Proof"], "#FFF3E0"),
        (8.0, 3.5, "Customer", ["Profiles", "Addresses", "Wishlist"], "#E8F0FE"),
        (3.5, 1.2, "Stock", ["Inventory", "Movements"], "#F3E8FF"),
        (6.5, 1.2, "Growth", ["Coupons", "Banners", "Reviews", "Recommendations"], "#FCE4EC"),
    ]
    for x, y, title, items, color in domains:
        _draw_box(ax, x, y, 2.3, 0.45, title, color, fs=8)
        for i, item in enumerate(items):
            _draw_box(ax, x, y - 0.55 - i * 0.45, 2.0, 0.35, item, "#FFFFFF", fs=6.5)
    ax.set_title("V1 Domain Grouping — Single Bag Store", fontsize=12, fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def make_v1_inventory_flow():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 3.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.2)
    ax.axis("off")
    steps = ["Customer", "Cart", "Inventory\n(Reserved)", "Order", "Inventory\n(OnHand -=)"]
    xs = [1, 2.8, 4.8, 6.8, 8.8]
    for x, name in zip(xs, steps):
        _draw_box(ax, x, 1.6, 1.5, 0.6, name, "#E8F0FE", fs=7)
    labels = ["Add item", "QtyReserved +=", "Checkout", "Sale movement"]
    for i in range(len(xs) - 1):
        _arrow(ax, xs[i] + 0.75, 1.6, xs[i + 1] - 0.75, 1.6)
        ax.text((xs[i] + xs[i + 1]) / 2, 2.35, labels[i], ha="center", fontsize=7, color="#4A5568")
    ax.set_title("V1 Inventory Flow (Single Warehouse)", fontsize=12, fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def make_v1_order_flow():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(11, 4))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 4)
    ax.axis("off")
    cod = ["Pending", "Processing", "Shipped", "Delivered", "Completed"]
    for i, s in enumerate(cod):
        _draw_box(ax, 0.8 + i * 2.0, 3.0, 1.5, 0.5, s, "#D4EDDA", fs=7)
        if i < len(cod) - 1:
            _arrow(ax, 0.8 + i * 2.0 + 0.75, 3.0, 0.8 + (i + 1) * 2.0 - 0.75, 3.0)
    ax.text(5.5, 3.65, "Cash On Delivery", ha="center", fontsize=9, fontweight="bold", color="#276749")
    wallet = ["Pending", "Awaiting\nPayment", "Proof\nSubmitted", "Confirmed", "Processing", "Shipped", "Delivered"]
    for i, s in enumerate(wallet):
        _draw_box(ax, 0.5 + i * 1.45, 1.15, 1.15, 0.55, s, "#FFF3E0", fs=6)
        if i < len(wallet) - 1:
            _arrow(ax, 0.5 + i * 1.45 + 0.58, 1.15, 0.5 + (i + 1) * 1.45 - 0.58, 1.15)
    ax.text(5.5, 1.85, "Wallet: InstaPay, Vodafone Cash, Orange, Etisalat, WE Pay", ha="center", fontsize=8,
            fontweight="bold", color="#C05621")
    ax.set_title("V1 Order & Payment Workflows", fontsize=12, fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def make_v1_simplification_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(9, 3.5))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 3.5)
    ax.axis("off")
    _draw_box(ax, 2.2, 1.8, 2.8, 1.4, "Full Architecture\n81 Tables\nEnterprise / Multi-warehouse", "#FED7D7", fs=9)
    _draw_box(ax, 6.8, 1.8, 2.8, 1.4, "V1 Commercial\n47 Tables\nSingle Bag Store", "#C6F6D5", fs=9)
    _arrow(ax, 3.6, 1.8, 5.4, 1.8)
    ax.text(4.5, 2.2, "42% reduction", ha="center", fontsize=10, fontweight="bold", color="#2D3748")
    ax.set_title("Architecture Simplification", fontsize=12, fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def add_diagram(doc, title, png_bytes, caption=None):
    add_para(doc, title, bold=True, size=12)
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.2))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption or f"Figure: {title}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title page
    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Cairo Bags")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("V1 Database Architecture")
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = s2.add_run("Commercial Release — Single Premium Bag Store")
    r3.font.size = Pt(14)
    r3.italic = True
    s3 = doc.add_paragraph()
    s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = s3.add_run("CODRA Development  |  " + datetime.now().strftime("%B %d, %Y"))
    r4.font.size = Pt(11)
    doc.add_page_break()

    # TOC
    add_heading(doc, "Table of Contents", 1)
    for item in [
        "1. Executive Summary",
        "2. V1 Design Principles & Enums",
        "3. Simplified Product Variants",
        "4. Shipping & Payments (Egypt)",
        "5. Updated ERD & Diagrams",
        "6. All V1 Table Definitions",
        "7. Final Table List (47 tables)",
        "8. Tables Removed from Full Architecture",
        "9. Why Each Table Exists",
        "10. Order & Payment Workflows",
        "11. Seed Data & V2 Scalability",
    ]:
        add_para(doc, item)
    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc,
             "Cairo Bags V1 is a production-ready database for a real single-store premium bag shop. "
             "It is NOT a marketplace, NOT multi-vendor, and NOT multi-warehouse.")
    add_diagram(doc, "Architecture Simplification Overview", make_v1_simplification_diagram())
    add_table(doc, ["Metric", "Full Design", "V1 Design"],
              [["Total tables", "81", "47"], ["Reduction", "—", "42% (~34 tables removed)"],
               ["Warehouses", "Multi", "Single implicit store"],
               ["Variant model", "Attribute engine", "ColorNameAr / ColorNameEn on variant"],
               ["Promotions", "Full engine", "Coupons only"],
               ["Collections", "Yes", "Removed"]])

    # 2. Principles
    add_heading(doc, "2. V1 Design Principles & Enums", 1)
    add_table(doc, ["Principle", "V1 Decision"],
              [["Store model", "Single bag shop, one stock pool"],
               ["Localization", "ProductTranslations, CategoryTranslations, bilingual banners"],
               ["Sellable unit", "ProductVariant (color + SKU + price)"],
               ["Stock", "Inventory 1:1 per variant + InventoryMovements"],
               ["Payments", "6 Egyptian methods; wallets need proof + admin approval"],
               ["Shipping", "Cairo / Giza / Other Governorates"],
               ["Audit", "CreatedAt, UpdatedAt, CreatedBy, UpdatedBy"],
               ["Money", "DECIMAL(18,2)  |  Timestamps: DATETIME2 UTC"]])

    add_heading(doc, "2.1 Core Enums", 2)
    add_table(doc, ["Enum", "Values"],
              [["OrderStatus", "Pending, AwaitingPayment, PaymentProofSubmitted, PaymentUnderReview, PaymentConfirmed, Processing, Shipped, Delivered, Completed, Cancelled, Refunded"],
               ["PaymentMethodType", "COD, InstaPay, VodafoneCash, OrangeCash, EtisalatCash, WEPay"],
               ["PaymentStatus", "Pending, ProofSubmitted, UnderReview, Confirmed, Rejected, Refunded"],
               ["ShippingZoneType", "Cairo, Giza, OtherGovernorates"],
               ["CouponType", "Percentage, FixedAmount, FreeShipping"],
               ["ReviewStatus", "Pending, Approved, Rejected"],
               ["NotificationType", "Order, Payment, Shipment, Promotion, Review, System"]])

    # 3. Variants
    add_heading(doc, "3. Simplified Product Variants", 1)
    add_para(doc, "Removed: AttributeDefinitions, AttributeOptions, VariantAttributeValues, ProductCategoryMappings")
    add_table(doc, ["Table", "Key Columns", "Purpose"],
              [["Products", "CategoryId, Status, IsFeatured, AverageRating", "Catalog core (non-translatable)"],
               ["ProductTranslations", "Name, Description, Slug (ar/en)", "Bilingual content + SEO"],
               ["ProductVariants", "ColorNameAr, ColorNameEn, Sku, Price, IsDefault", "Sellable SKU — color only"],
               ["ProductImages", "ProductId, VariantId, ImageUrl, IsPrimary", "Gallery + per-color images"],
               ["Inventory", "VariantId (unique), QuantityOnHand, QuantityReserved, RowVersion", "Single-store stock"]])

    add_heading(doc, "3.1 ProductVariant Example", 2)
    add_table(doc, ["Column", "Example"],
              [["ColorNameAr", "أسود"], ["ColorNameEn", "Black"], ["Sku", "CB-TOTE-BLK"],
               ["Price", "1,250.00 EGP"], ["IsDefault", "true"], ["Stock (Inventory)", "QuantityOnHand = 15"]])

    # 4. Shipping & Payments
    add_heading(doc, "4. Shipping & Payments (Egypt)", 1)
    add_heading(doc, "4.1 Shipping Zones", 2)
    add_table(doc, ["Zone", "Code", "Example Governorates"],
              [["Cairo", "CAIRO", "Cairo"], ["Giza", "GIZA", "Giza"],
               ["Other Governorates", "OTHER", "Alexandria, Dakahlia, Sharqia, all others"]])
    add_para(doc, "ShippingRates merged into ShippingZones: BaseShippingFee, FreeShippingThreshold per zone.")

    add_heading(doc, "4.2 Payment Methods", 2)
    add_table(doc, ["#", "Code", "Method", "Proof Required", "Admin Approval"],
              [["1", "COD", "Cash On Delivery", "No", "No"],
               ["2", "INSTAPAY", "InstaPay", "Yes", "Yes"],
               ["3", "VODAFONE_CASH", "Vodafone Cash", "Yes", "Yes"],
               ["4", "ORANGE_CASH", "Orange Cash", "Yes", "Yes"],
               ["5", "ETISALAT_CASH", "Etisalat Cash", "Yes", "Yes"],
               ["6", "WE_PAY", "WE Pay", "Yes", "Yes"]])

    # 5. Diagrams
    add_heading(doc, "5. Updated ERD & Diagrams", 1)
    add_diagram(doc, "V1 Module Map", make_v1_module_diagram())
    add_diagram(doc, "V1 Domain Grouping", make_v1_domain_diagram())
    add_diagram(doc, "V1 Core ERD", make_v1_erd_diagram())
    add_diagram(doc, "V1 Inventory Flow", make_v1_inventory_flow())
    add_diagram(doc, "V1 Order & Payment Workflows", make_v1_order_flow())

    doc.add_page_break()

    # 6. Table definitions by domain
    add_heading(doc, "6. All V1 Table Definitions", 1)

    domains = [
        ("A. Identity & Access (9 tables)", [
            ("AspNetUsers + 6 Identity tables", "Authentication, Admin/Customer roles, JWT refresh, Google login"),
            ("CustomerProfiles", "UserId (unique), FirstName, LastName, DisplayName, ProfileImageUrl, MarketingOptIn"),
            ("PasswordResetOtps", "Hashed OTP, NormalizedEmail, ExpiresAtUtc"),
        ]),
        ("B. Reference (3 tables)", [
            ("Languages", "en, ar — RTL flag for Arabic"),
            ("Currencies", "EGP default"),
            ("Governorates", "27 governorates mapped to ShippingZoneId"),
        ]),
        ("C. Catalog (6 tables)", [
            ("Categories + CategoryTranslations", "Hierarchy, bilingual Name/Slug/Description"),
            ("Products + ProductTranslations", "Status, featured flags, bilingual Name/Description/Slug"),
            ("ProductVariants", "ColorNameAr, ColorNameEn, Sku, Price, IsDefault, Status"),
            ("ProductImages", "Gallery, variant-specific images, AltText ar/en"),
        ]),
        ("D. Inventory (2 tables)", [
            ("Inventory", "VariantId unique, QuantityOnHand, QuantityReserved, LowStockThreshold, RowVersion"),
            ("InventoryMovements", "Sale, Return, Adjustment, Reservation — full audit trail"),
        ]),
        ("E. Cart & Wishlist (4 tables)", [
            ("Carts + CartItems", "User or guest session, variant + quantity + unit price"),
            ("Wishlists + WishlistItems", "One wishlist per user, saved products/variants"),
        ]),
        ("F. Orders (4 tables)", [
            ("Orders", "OrderNumber, status, totals, shipping snapshot on order, tracking fields"),
            ("OrderItems", "Snapshot: Sku, ProductName ar/en, ColorName ar/en, price, image"),
            ("OrderStatusHistories", "Full status transition audit"),
            ("ShippingAddresses", "Customer address book"),
        ]),
        ("G. Shipping (2 tables)", [
            ("ShippingZones", "Cairo/Giza/Other, BaseShippingFee, FreeShippingThreshold"),
            ("ShippingZoneTranslations", "Bilingual zone names"),
        ]),
        ("H. Payments (4 tables)", [
            ("PaymentMethods + Translations", "6 methods, bilingual checkout instructions"),
            ("OrderPayments", "Status, amount, sender wallet, admin review fields"),
            ("PaymentProofImages", "Transfer screenshot uploads"),
        ]),
        ("I. Reviews (2 tables)", [
            ("ProductReviews", "Rating 1-5, verified purchase, moderation status, admin response"),
            ("ReviewImages", "Customer review photos"),
        ]),
        ("J. Coupons (2 tables)", [
            ("Coupons", "Code, type, scope via nullable CategoryId/ProductId, usage limits"),
            ("CouponUsages", "Per-user and per-order tracking"),
        ]),
        ("K. Notifications (1 table)", [
            ("Notifications", "In-app inbox: Order, Payment, Shipment, Promotion alerts"),
        ]),
        ("L. Banners (1 table)", [
            ("Banners", "Inline bilingual Title/Subtitle/CTA/Images, Placement, scheduling"),
        ]),
        ("M. Analytics (2 tables)", [
            ("DailySalesSummaries", "Revenue, orders, AOV per day"),
            ("ProductPerformanceMetrics", "Views, conversion, revenue per product"),
        ]),
        ("N. Recommendations (3 tables)", [
            ("ProductRecommendations", "Manual + computed cross-sell"),
            ("UserProductViews", "Browse behavior"),
            ("TrendingProducts", "Precomputed homepage trending"),
        ]),
        ("O. System (2 tables)", [
            ("SystemSettings", "Store name ar/en, maintenance, min order, contact info"),
            ("AuditLogs", "Admin change accountability (JSON old/new values)"),
        ]),
    ]
    for domain_name, tables in domains:
        add_heading(doc, domain_name, 2)
        add_table(doc, ["Table", "Description"], tables)

    # 7. Final table list
    add_heading(doc, "7. Final Table List — 47 Tables", 1)
    final_tables = [
        (1, "AspNetUsers", "Identity"), (2, "AspNetRoles", "Identity"),
        (3, "AspNetUserRoles", "Identity"), (4, "AspNetUserClaims", "Identity"),
        (5, "AspNetRoleClaims", "Identity"), (6, "AspNetUserLogins", "Identity"),
        (7, "AspNetUserTokens", "Identity"), (8, "CustomerProfiles", "Customer"),
        (9, "PasswordResetOtps", "Auth"), (10, "Languages", "Reference"),
        (11, "Currencies", "Reference"), (12, "Governorates", "Reference"),
        (13, "Categories", "Catalog"), (14, "CategoryTranslations", "Catalog"),
        (15, "Products", "Catalog"), (16, "ProductTranslations", "Catalog"),
        (17, "ProductVariants", "Catalog"), (18, "ProductImages", "Catalog"),
        (19, "Inventory", "Inventory"), (20, "InventoryMovements", "Inventory"),
        (21, "Carts", "Commerce"), (22, "CartItems", "Commerce"),
        (23, "Wishlists", "Commerce"), (24, "WishlistItems", "Commerce"),
        (25, "Orders", "Orders"), (26, "OrderItems", "Orders"),
        (27, "OrderStatusHistories", "Orders"), (28, "ShippingAddresses", "Orders"),
        (29, "ShippingZones", "Shipping"), (30, "ShippingZoneTranslations", "Shipping"),
        (31, "PaymentMethods", "Payments"), (32, "PaymentMethodTranslations", "Payments"),
        (33, "OrderPayments", "Payments"), (34, "PaymentProofImages", "Payments"),
        (35, "ProductReviews", "Reviews"), (36, "ReviewImages", "Reviews"),
        (37, "Coupons", "Marketing"), (38, "CouponUsages", "Marketing"),
        (39, "Notifications", "Notifications"), (40, "Banners", "Marketing"),
        (41, "DailySalesSummaries", "Analytics"), (42, "ProductPerformanceMetrics", "Analytics"),
        (43, "ProductRecommendations", "Recommendations"), (44, "UserProductViews", "Recommendations"),
        (45, "TrendingProducts", "Recommendations"), (46, "SystemSettings", "System"),
        (47, "AuditLogs", "System"),
    ]
    add_table(doc, ["#", "Table", "Domain"], final_tables)

    add_heading(doc, "7.1 Table Count by Category", 2)
    add_table(doc, ["Category", "Count"],
              [["ASP.NET Identity", "7"], ["Auth & Customer", "2"], ["Reference", "3"],
               ["Catalog", "6"], ["Inventory", "2"], ["Cart & Wishlist", "4"],
               ["Orders & Addresses", "4"], ["Shipping", "2"], ["Payments", "4"],
               ["Reviews", "2"], ["Coupons", "2"], ["Notifications", "1"],
               ["Banners", "1"], ["Analytics", "2"], ["Recommendations", "3"],
               ["System", "2"], ["TOTAL V1", "47"]])

    # 8. Removed
    add_heading(doc, "8. Tables Removed from Full Architecture", 1)
    removed = [
        ("Collections, CollectionTranslations, CollectionProducts", "Not needed for V1 catalog"),
        ("Warehouses", "Single implicit stock pool"),
        ("AttributeDefinitions + Options + VariantAttributeValues (5 tables)", "Color on variant directly"),
        ("ProductCategoryMappings", "One CategoryId on product"),
        ("StockReservations", "Inventory.QuantityReserved handles V1"),
        ("OrderAddresses", "Shipping snapshot merged into Orders"),
        ("ShippingRates", "Merged into ShippingZones"),
        ("Shipments, ShipmentTrackingEvents", "Tracking on Orders"),
        ("AdminProfiles", "Admin uses Identity roles only"),
        ("Promotions module (5 tables)", "Coupons only in V1"),
        ("CouponScopes", "Nullable CategoryId/ProductId on coupon"),
        ("ReviewHelpfulVotes", "Not required V1"),
        ("BannerPlacements, BannerTranslations", "Inline bilingual on Banners"),
        ("NotificationTemplates, UserNotificationPreferences", "Single Notifications table"),
        ("RecommendationRules, UserProductInteractions", "Simplified recommendation engine"),
        ("AnalyticsEvents, Category/Payment/Shipping analytics", "Pre-aggregated tables only"),
        ("Loyalty program fields", "Removed from CustomerProfiles"),
    ]
    add_table(doc, ["Removed", "Reason"], removed)
    add_para(doc, "Total removed: ~34 tables from the original 81-table design.", bold=True)

    # 9. Why each table exists
    add_heading(doc, "9. Why Each Remaining Table Exists", 1)
    why = [
        ("AspNetUsers (+ Identity)", "Production authentication for real customers and admins"),
        ("CustomerProfiles", "Profile data separate from login credentials"),
        ("PasswordResetOtps", "Secure password recovery"),
        ("Languages / Currencies / Governorates", "Arabic-English support and Egyptian addresses"),
        ("Categories + Translations", "Bag taxonomy with bilingual SEO"),
        ("Products + Translations", "Catalog with bilingual names and descriptions"),
        ("ProductVariants", "What actually goes in the cart (color + SKU + price)"),
        ("ProductImages", "Premium product gallery and per-color images"),
        ("Inventory", "Prevent overselling — single store stock"),
        ("InventoryMovements", "Accountability for every stock change"),
        ("Carts + CartItems", "Shopping before checkout"),
        ("Wishlists + WishlistItems", "Save bags; feeds recommendations"),
        ("Orders + OrderItems", "Purchase records with immutable snapshots"),
        ("OrderStatusHistories", "Customer tracking and dispute resolution"),
        ("ShippingAddresses", "Repeat customer address book"),
        ("ShippingZones + Translations", "Cairo / Giza / Other pricing"),
        ("PaymentMethods + Translations", "6 Egyptian methods with bilingual instructions"),
        ("OrderPayments + ProofImages", "Wallet payment verification workflow"),
        ("ProductReviews + ReviewImages", "Social proof for premium bags"),
        ("Coupons + CouponUsages", "Discount codes with limits"),
        ("Notifications", "Order and payment status alerts"),
        ("Banners", "Real store homepage marketing"),
        ("DailySalesSummaries", "Admin revenue dashboard"),
        ("ProductPerformanceMetrics", "Best sellers and conversion data"),
        ("ProductRecommendations", "Cross-sell on product pages"),
        ("UserProductViews / TrendingProducts", "Personalization and homepage trending"),
        ("SystemSettings", "Store configuration and maintenance mode"),
        ("AuditLogs", "Admin accountability on prices and orders"),
    ]
    add_table(doc, ["Table(s)", "Why"], why)

    # 10. Workflows
    add_heading(doc, "10. Order & Payment Workflows", 1)
    add_table(doc, ["Method", "Flow"],
              [["Cash On Delivery", "Pending → Processing → Shipped → Delivered → Completed"],
               ["Wallet methods", "Pending → AwaitingPayment → PaymentProofSubmitted → PaymentUnderReview → PaymentConfirmed → Processing → Shipped → Delivered"]])

    # 11. Seed & V2
    add_heading(doc, "11. Seed Data & V2 Scalability", 1)
    add_heading(doc, "11.1 Required Seed Data", 2)
    add_table(doc, ["Table", "Seed"],
              [["Languages", "en, ar"], ["Currencies", "EGP"],
               ["AspNetRoles", "Admin, Customer"], ["ShippingZones", "Cairo, Giza, Other"],
               ["Governorates", "27 Egyptian governorates"], ["PaymentMethods", "6 methods"],
               ["SystemSettings", "Cairo Bags default row"]])

    add_heading(doc, "11.2 V2 Path (without over-engineering V1)", 2)
    add_table(doc, ["Future Need", "V2 Addition"],
              [["Collections / lookbooks", "Add 3 tables"], ["Loyalty program", "LoyaltyTransactions table"],
               ["Multi-warehouse", "Warehouses + WarehouseId on Inventory"],
               ["Rich variant attributes", "Reintroduce attribute tables if needed"],
               ["Promotion engine", "Promotions module"], ["Raw analytics", "AnalyticsEvents table"]])

    doc.add_page_break()
    add_heading(doc, "Document Information", 1)
    add_para(doc, "Cairo Bags V1 Database Architecture — CODRA Development")
    add_para(doc, "Architecture documentation only. No code, entities, DbContext, or migrations.")
    add_para(doc, f"File: {OUTPUT}")
    add_para(doc, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    return doc


if __name__ == "__main__":
    print("Building Cairo Bags V1 Database Architecture report...")
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Size: {os.path.getsize(OUTPUT):,} bytes")
