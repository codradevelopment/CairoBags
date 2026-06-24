# -*- coding: utf-8 -*-
"""Generate Cairo Bags Frontend Architecture & Roadmap Word report on Desktop."""
import io
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")
OUTPUT = os.path.join(DESKTOP, "Cairo_Bags_Frontend_Architecture_Roadmap.docx")

GOLD = "C9A962"
BLACK = "0A0A0A"
HEADER = "1A1A1A"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def add_table(doc, headers, rows, header_color=HEADER):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph()


def make_frontend_layers_diagram():
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")

    def box(x, y, w, h, text, color="#FAFAFA", edge="#0A0A0A", fs=8):
        patch = FancyBboxPatch(
            (x - w / 2, y - h / 2), w, h,
            boxstyle="round,pad=0.02", facecolor=color, edgecolor=edge, linewidth=1.2)
        ax.add_patch(patch)
        ax.text(x, y, text, ha="center", va="center", fontsize=fs, fontweight="bold", color="#0A0A0A")

    box(5, 5.2, 3.0, 0.55, "React Pages (Store + Admin + Auth)", "#FFF8E7", fs=9)
    box(5, 4.2, 3.2, 0.55, "Layouts + Reusable Components", "#FAFAFA")
    box(5, 3.2, 3.2, 0.55, "Context API (Auth, Cart, Locale, Notifications)", "#F5F5F5")
    box(5, 2.2, 3.2, 0.55, "Hooks + Validators + Formatters", "#F0F0F0")
    box(5, 1.2, 3.2, 0.55, "Axios API Layer", "#E8E8E8")
    box(5, 0.35, 3.2, 0.55, "Cairo Bags Backend V1 (Frozen)", "#C9A962", edge="#0A0A0A", fs=9)

    for y1, y2 in [(5.0, 4.5), (4.0, 3.5), (3.0, 2.5), (2.0, 1.5), (1.0, 0.65)]:
        ax.annotate("", xy=(5, y2), xytext=(5, y1),
                    arrowprops=dict(arrowstyle="->", color="#C9A962", lw=1.5))

    ax.set_title("Cairo Bags Frontend Architecture Layers", fontsize=13, fontweight="bold", color="#0A0A0A")
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def make_roadmap_diagram():
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.5)
    ax.axis("off")

    phases = [
        (1.0, "Phase 0\nFoundation", "#FAFAFA"),
        (3.0, "Phase 1\nStorefront", "#FFF8E7"),
        (5.0, "Phase 2\nCheckout", "#F5F0E0"),
        (7.0, "Phase 3\nAdmin", "#F0EBE0"),
        (9.0, "Phase 4\nPolish", "#C9A962"),
    ]
    for x, label, color in phases:
        patch = FancyBboxPatch(
            (x - 0.85, 1.5), 1.7, 1.2,
            boxstyle="round,pad=0.03", facecolor=color, edgecolor="#0A0A0A", linewidth=1.2)
        ax.add_patch(patch)
        ax.text(x, 2.1, label, ha="center", va="center", fontsize=8, fontweight="bold", color="#0A0A0A")

    for x1, x2 in [(1.85, 2.15), (3.85, 4.15), (5.85, 6.15), (7.85, 8.15)]:
        ax.annotate("", xy=(x2, 2.1), xytext=(x1, 2.1),
                    arrowprops=dict(arrowstyle="->", color="#0A0A0A", lw=1.2))

    ax.text(5, 3.8, "7 Weeks to Frontend V1  |  Post-V1: Coupons, Banners, Wishlist, Reviews",
            ha="center", fontsize=9, color="#555555")
    ax.set_title("Frontend Development Roadmap", fontsize=13, fontweight="bold", color="#0A0A0A")

    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def build_document():
    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Cairo Bags")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("React Frontend Architecture & Development Roadmap")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x62)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Backend V1 Complete & Frozen  |  Architecture Only — No Code")
    run.font.size = Pt(12)
    run.font.italic = True

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub3.add_run(datetime.now().strftime("%B %Y"))
    run.font.size = Pt(11)

    doc.add_page_break()

    # Executive summary
    add_heading(doc, "Executive Summary", 1)
    add_para(doc,
             "This document defines the complete React frontend architecture for Cairo Bags — "
             "a premium luxury bag e-commerce platform with Arabic/English support, "
             "customer storefront, and admin panel.")
    add_para(doc, "Technology stack: React, React Router, Axios, Context API, Tailwind CSS.")
    add_table(doc, ["Area", "Decision"], [
        ["Stack", "React + Vite + TypeScript recommended"],
        ["State", "Context API (Auth, Cart, Locale, Notifications)"],
        ["API", "Single Axios client mirroring frozen Backend V1"],
        ["Brand", "Black / White / Gold — Premium, Luxury, Modern"],
        ["Timeline", "~7 weeks to Frontend V1"],
    ])

    # Brand
    add_heading(doc, "Brand & Design System", 1)
    add_table(doc, ["Token", "Hex", "Usage"], [
        ["Black", "#0A0A0A", "Headers, primary text, admin sidebar"],
        ["White", "#FAFAFA", "Backgrounds, cards"],
        ["Gold", "#C9A962", "CTAs, badges, dividers, hover"],
        ["Gold Muted", "#E8D5A3", "Borders, subtle highlights"],
    ])
    add_para(doc, "Typography: Serif headings (Playfair/Cormorant) + Sans body (Inter/DM Sans). Arabic: Tajawal or Cairo.")
    add_para(doc, "Style: Premium, Luxury, Modern, Elegant — generous whitespace, restrained gold accents.")

    doc.add_page_break()

    # Diagram
    add_heading(doc, "Architecture Overview", 1)
    doc.add_picture(io.BytesIO(make_frontend_layers_diagram()), width=Inches(6.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 1. Folder Structure
    add_heading(doc, "1. Folder Structure", 1)
    add_code_block(doc, """cairo-bags-web/
├── public/
├── src/
│   ├── app/                 # App, router, providers
│   ├── assets/
│   ├── config/              # env, constants, routes
│   ├── api/                 # Axios modules (mirror backend)
│   ├── contexts/            # Auth, Cart, Locale, Notifications
│   ├── hooks/
│   ├── layouts/             # Store, Auth, Account, Admin, Minimal
│   ├── pages/
│   │   ├── store/
│   │   ├── auth/
│   │   ├── account/
│   │   └── admin/
│   ├── components/
│   │   ├── ui/              # Design system primitives
│   │   ├── store/
│   │   ├── account/
│   │   ├── admin/
│   │   └── shared/
│   ├── lib/                 # i18n, signalr, storage, formatters
│   ├── styles/
│   └── types/
├── tailwind.config.js
└── vite.config.ts""")

    add_para(doc, "Conventions:", bold=True)
    add_bullet(doc, "One folder per page; co-locate page-specific subcomponents.")
    add_bullet(doc, "API modules map 1:1 to backend controllers.")
    add_bullet(doc, "Admin pages never import store pages; shared logic in hooks/lib.")

    doc.add_page_break()

    # 2. Routing
    add_heading(doc, "2. Routing Structure", 1)
    add_heading(doc, "Public Store Routes", 2)
    add_table(doc, ["Path", "Page", "Layout", "Auth"], [
        ["/", "Home", "StoreLayout", "—"],
        ["/shop", "Shop", "StoreLayout", "—"],
        ["/shop/:categorySlug", "Shop (filtered)", "StoreLayout", "—"],
        ["/categories", "Categories", "StoreLayout", "—"],
        ["/product/:slug", "Product Details", "StoreLayout", "—"],
        ["/cart", "Cart", "StoreLayout", "—"],
        ["/login", "Login", "AuthLayout", "Guest only"],
        ["/register", "Register", "AuthLayout", "Guest only"],
        ["/forgot-password", "Forgot Password", "AuthLayout", "Guest only"],
    ])

    add_heading(doc, "Protected Store Routes", 2)
    add_table(doc, ["Path", "Page", "Layout", "Auth"], [
        ["/checkout", "Checkout", "MinimalLayout", "Customer"],
        ["/order-success/:orderId", "Order Success", "MinimalLayout", "Customer"],
        ["/account", "Profile", "AccountLayout", "Customer"],
        ["/account/orders", "My Orders", "AccountLayout", "Customer"],
        ["/account/orders/:id", "Order Details", "AccountLayout", "Customer"],
        ["/account/orders/:id/payment", "Payment Proof", "AccountLayout", "Customer"],
    ])

    add_heading(doc, "Admin Routes (/admin/*)", 2)
    add_table(doc, ["Path", "Page", "Role"], [
        ["/admin", "Dashboard", "Admin"],
        ["/admin/categories", "Categories CRUD", "Admin"],
        ["/admin/products", "Products list", "Admin"],
        ["/admin/products/:id", "Product edit", "Admin"],
        ["/admin/inventory", "Inventory list", "Admin"],
        ["/admin/inventory/:variantId", "Inventory detail", "Admin"],
        ["/admin/orders", "Orders + filters", "Admin"],
        ["/admin/orders/:id", "Order actions", "Admin"],
        ["/admin/payments", "Pending reviews", "Admin"],
        ["/admin/payments/:id", "Approve/Reject", "Admin"],
        ["/admin/coupons", "Coupons", "Admin"],
        ["/admin/banners", "Banners", "Admin"],
        ["/admin/settings", "Settings", "Admin"],
    ])

    add_para(doc, "Route guards: GuestRoute, ProtectedRoute, AdminRoute.")
    add_para(doc, "Locale URL prefix recommended: /ar/shop and /en/shop for SEO and RTL/LTR.")

    doc.add_page_break()

    # 3. Layouts
    add_heading(doc, "3. Layout Architecture", 1)
    add_table(doc, ["Layout", "Used For", "Key Elements"], [
        ["StoreLayout", "Browse, shop, cart", "Header, footer, mobile bottom nav"],
        ["AuthLayout", "Login, register", "Centered card, minimal branding"],
        ["AccountLayout", "Profile, orders", "Side nav (desktop) / tabs (mobile)"],
        ["MinimalLayout", "Checkout, success", "Logo + stepper, no distractions"],
        ["AdminLayout", "All admin", "Dark sidebar, top bar, notifications"],
    ])
    add_para(doc, "Breakpoints: sm 640 | md 768 | lg 1024 | xl 1280 (max-w-7xl store container).")

    doc.add_page_break()

    # 4. Context
    add_heading(doc, "4. Context Architecture", 1)
    add_code_block(doc, """providers.tsx
├── LocaleProvider       # lang, dir, t()
├── AuthProvider         # user, tokens, login/logout/refresh
├── CartProvider         # cart, sessionId, merge on login
└── NotificationProvider # SignalR + unread count + toast""")

    add_heading(doc, "AuthContext", 2)
    add_bullet(doc, "State: user, roles, accessToken, refreshToken")
    add_bullet(doc, "Actions: login, register, logout, refreshToken, updateProfile")
    add_bullet(doc, "Flow: Login → tokens → GET /api/Account/Me → merge guest cart")

    add_heading(doc, "CartContext", 2)
    add_bullet(doc, "GET /api/cart with X-Session-Id for guests")
    add_bullet(doc, "Reconcile stockChanged flag from API")

    add_heading(doc, "LocaleContext", 2)
    add_bullet(doc, "ar = RTL, en = LTR; persist in localStorage")
    add_bullet(doc, "Update document.documentElement.dir and lang")

    add_heading(doc, "NotificationContext", 2)
    add_bullet(doc, "REST + SignalR ReceiveNotification hub")
    add_bullet(doc, "Connect after login; disconnect on logout")

    add_para(doc, "Rule: Contexts orchestrate — no raw Axios inside Context (use api/ modules).")

    doc.add_page_break()

    # 5. API Layer
    add_heading(doc, "5. API Layer Structure", 1)
    add_table(doc, ["Module", "Backend Endpoints"], [
        ["auth.api", "register, LogIn, refresh-token, Me, forgot-password"],
        ["catalog.api", "categories, products, search, featured, images"],
        ["cart.api", "cart CRUD, merge"],
        ["checkout.api", "POST /api/checkout"],
        ["orders.api", "my orders, detail, cancel"],
        ["payments.api", "proof upload, payment status"],
        ["recommendations.api", "trending, recently-viewed, similar, bought-together"],
        ["notifications.api", "list, read, unread-count"],
        ["admin/*.api", "admin CRUD + order/payment actions"],
    ])

    add_para(doc, "Axios interceptors:", bold=True)
    add_bullet(doc, "Request: Authorization Bearer + X-Session-Id")
    add_bullet(doc, "Response 401: refresh once → retry or logout")
    add_bullet(doc, "Response 409: surface concurrency_conflict (inventory/checkout)")

    doc.add_page_break()

    # 6. Components
    add_heading(doc, "6. Reusable Components Structure", 1)
    add_heading(doc, "components/ui/ — Design System", 2)
    add_para(doc, "Button, Input, Card, Badge, Modal, Drawer, Spinner, Skeleton, Toast, "
             "EmptyState, Pagination, Tabs, Stepper, Image, Price, StockBadge")

    add_heading(doc, "components/store/", 2)
    add_para(doc, "ProductCard, ProductGallery, VariantPicker, CategoryCard, CartItemRow, "
             "CartSummary, ShippingAddressForm, PaymentMethodSelector, CouponInput, "
             "PaymentProofUpload, OrderStatusTimeline, RecommendationRow, Header, Footer, MobileBottomNav")

    add_heading(doc, "components/admin/", 2)
    add_para(doc, "AdminSidebar, DataTable, StatusBadge, ConfirmDialog, ImageUploader, "
             "InventoryAdjustForm, OrderActionPanel, PaymentReviewPanel, FilterBar")

    add_heading(doc, "components/shared/", 2)
    add_para(doc, "ErrorBoundary, ProtectedRoute, AdminRoute, GuestRoute, SeoHead, "
             "LanguageSwitcher, NotificationBell, Breadcrumbs")

    doc.add_page_break()

    # 7. Roadmap
    add_heading(doc, "7. Frontend Development Roadmap", 1)
    doc.add_picture(io.BytesIO(make_roadmap_diagram()), width=Inches(6.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Phase 0 — Foundation (Week 1)", 2)
    add_table(doc, ["Task", "Deliverable"], [
        ["Vite + React + TS + Tailwind", "Project scaffold"],
        ["Design tokens (black/white/gold)", "tailwind.config.js"],
        ["i18n skeleton", "LocaleContext + ar/en JSON"],
        ["Axios client + interceptors", "api/client.ts"],
        ["Router + layouts shell", "Empty pages, guards"],
        ["AuthContext + token refresh", "Login flow wired"],
    ])
    add_para(doc, "Exit: App boots, login works, locale switches RTL/LTR.")

    add_heading(doc, "Phase 1 — Storefront Core (Weeks 2–3)", 2)
    add_table(doc, ["Sprint", "Scope"], [
        ["1.1 Catalog", "Home, Shop, Categories, Product Details + recommendations prep"],
        ["1.2 Cart", "CartContext, Cart page, guest sessionId, merge on login"],
        ["1.3 Recommendations", "Trending, similar, bought-together on Home/Product"],
    ])
    add_para(doc, "Exit: Browse → add to cart (guest + auth).")

    add_heading(doc, "Phase 2 — Checkout & Orders (Week 4)", 2)
    add_table(doc, ["Sprint", "Scope"], [
        ["2.1 Auth UX", "Register, Forgot Password, Profile"],
        ["2.2 Checkout", "Address, payment method, coupon, Order Success"],
        ["2.3 My Orders", "List, detail, cancel"],
        ["2.4 Wallet payments", "Payment proof upload"],
    ])
    add_para(doc, "Exit: Full purchase flow COD + wallet proof submission.")

    add_heading(doc, "Phase 3 — Admin Core (Weeks 5–6)", 2)
    add_table(doc, ["Sprint", "Scope"], [
        ["3.1 Admin shell", "AdminLayout, Dashboard KPIs"],
        ["3.2 Catalog admin", "Categories, Products, image upload"],
        ["3.3 Inventory", "List, low stock, adjust, movements"],
        ["3.4 Orders admin", "List, filters, status transitions"],
        ["3.5 Payments admin", "Pending queue, approve/reject"],
    ])
    add_para(doc, "Exit: Admin manages catalog, fulfills orders, reviews payments.")

    add_heading(doc, "Phase 4 — Polish & Realtime (Week 7)", 2)
    add_table(doc, ["Sprint", "Scope"], [
        ["4.1 Notifications", "SignalR hub, bell, toasts"],
        ["4.2 Settings admin", "Maintenance mode, beta flags"],
        ["4.3 UX polish", "Skeletons, empty states, 404"],
        ["4.4 SEO & perf", "Meta tags, lazy images, code split"],
    ])

    add_heading(doc, "Phase 5 — Post-V1 (Deferred)", 2)
    add_table(doc, ["Feature", "Depends On"], [
        ["Coupons admin + checkout UI", "Coupons API sprint"],
        ["Banners admin + Home hero", "Banners API sprint"],
        ["Reviews & ratings", "Reviews API"],
        ["Wishlist page", "Wishlist API"],
        ["Address book CRUD", "Shipping addresses API"],
        ["Dashboard analytics", "Analytics API"],
    ])

    doc.add_page_break()

    add_heading(doc, "Cross-Cutting Checklist", 1)
    checklist = [
        "Arabic + English for all user-facing strings",
        "RTL layout tested on every new page",
        "Mobile-first before desktop enhancements",
        "Loading + error + empty states on every fetch",
        "returnUrl preserved through login redirect",
        "EGP currency formatting (1,200.00 EGP)",
        "Image URLs prefixed with API/static host",
        "Admin destructive actions require confirmation",
    ]
    for item in checklist:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "Document Information", 1)
    add_para(doc, "Cairo Bags — React Frontend Architecture & Roadmap")
    add_para(doc, "Architecture documentation only. No React code included.")
    add_para(doc, f"Output file: {OUTPUT}")
    add_para(doc, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    return doc


if __name__ == "__main__":
    print("Building Cairo Bags Frontend Architecture report...")
    document = build_document()
    document.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Size: {os.path.getsize(OUTPUT):,} bytes")
