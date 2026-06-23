# -*- coding: utf-8 -*-
"""Generate Cairo Bags V1 FINAL Database Architecture Word report on Desktop."""
import io
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")
OUTPUT = os.path.join(DESKTOP, "Cairo_Bags_V1_Final_Database_Architecture.docx")


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def add_table(doc, headers, rows, header_color="1A365D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def _draw_box(ax, x, y, w, h, text, color="#E8F0FE", edge="#1A365D", fs=8):
    from matplotlib.patches import FancyBboxPatch
    box = FancyBboxPatch((x - w / 2, y - h / 2), w, h, boxstyle="round,pad=0.02",
                         facecolor=color, edgecolor=edge, linewidth=1.2)
    ax.add_patch(box)
    ax.text(x, y, text, ha="center", va="center", fontsize=fs, fontweight="bold",
            color="#1A365D", wrap=True)


def _arrow(ax, x1, y1, x2, y2):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="->", color="#4A5568", lw=1.2))


def _save_fig(fig):
    buf = io.BytesIO()
    import matplotlib.pyplot as plt
    plt.tight_layout()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def make_final_module_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    groups = [
        (1.8, 4.8, "Identity", ["Auth + Roles", "Profiles", "Password Reset"], "#E8F0FE"),
        (5.0, 4.8, "Catalog", ["Products", "Variants", "Translations", "Images"], "#D4EDDA"),
        (8.2, 4.8, "Commerce", ["Cart", "Orders", "Payments + Proof"], "#FFF3E0"),
        (2.0, 2.5, "Operations", ["Inventory", "Notifications", "Audit Logs"], "#F3E8FF"),
        (5.0, 2.5, "Fulfillment", ["Governorates", "Shipping Zones", "Addresses"], "#E0F7FA"),
        (8.0, 2.5, "Marketing", ["Coupons", "Banners", "Reviews + Images"], "#FCE4EC"),
        (5.0, 0.7, "Insights (Lean)", ["DailySalesSummaries", "UserProductViews", "TrendingProducts"], "#C6F6D5"),
    ]
    for x, y, title, items, color in groups:
        _draw_box(ax, x, y + 0.4, 2.1, 0.42, title, color, fs=9)
        for i, item in enumerate(items):
            _draw_box(ax, x, y - 0.15 - i * 0.42, 1.85, 0.34, item, "#FFFFFF", fs=6.5)
    ax.set_title("Cairo Bags V1 FINAL — Module Structure (43 Tables)", fontsize=12, fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def make_final_erd_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(11, 7.5))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 7.5)
    ax.axis("off")
    entities = [
        (1.2, 6.5, "Users"), (3.2, 6.5, "Profiles"),
        (1.2, 5.2, "ShipAddr"), (5.0, 6.5, "Categories"),
        (7.0, 6.5, "Products"), (9.2, 6.5, "Variants"),
        (7.0, 5.2, "Translations"), (9.2, 5.2, "Inventory"),
        (2.5, 3.7, "Carts"), (4.5, 3.7, "CartItems"),
        (6.5, 3.7, "Orders"), (8.5, 3.7, "OrderItems"),
        (6.5, 2.5, "Payments"), (8.5, 2.5, "ProofImages"),
        (2.5, 2.5, "Wishlists"), (4.5, 2.5, "Reviews"),
        (1.2, 1.1, "Governorates"), (3.5, 1.1, "ShipZones"),
        (5.5, 1.1, "Coupons"), (7.5, 1.1, "Views"),
        (9.2, 1.1, "Trending"),
    ]
    for x, y, name in entities:
        _draw_box(ax, x, y, 1.45, 0.4, name, "#E8F0FE", fs=6.5)
    pairs = [
        ((1.2, 6.28), (3.2, 6.28)), ((5.0, 6.28), (7.0, 6.28)), ((7.0, 6.28), (9.2, 6.28)),
        ((9.2, 6.28), (9.2, 5.42)), ((7.0, 4.98), (7.0, 5.42)),
        ((2.5, 3.48), (4.5, 3.48)), ((6.5, 3.48), (8.5, 3.48)),
        ((6.5, 3.28), (6.5, 2.72)), ((8.5, 3.28), (8.5, 2.72)),
        ((1.2, 1.32), (3.5, 1.32)), ((7.5, 1.32), (9.2, 1.32)),
    ]
    for a, b in pairs:
        _arrow(ax, *a, *b)
    ax.set_title("V1 FINAL ERD — 43 Tables", fontsize=12, fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def make_insights_pipeline_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 3.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.2)
    ax.axis("off")
    _draw_box(ax, 1.5, 1.6, 2.0, 0.7, "UserProductViews\n(browse events)", "#E8F0FE", fs=7)
    _draw_box(ax, 5.0, 1.6, 2.0, 0.7, "Nightly Job", "#FFF3E0", fs=8)
    _draw_box(ax, 8.5, 2.2, 1.8, 0.55, "TrendingProducts", "#C6F6D5", fs=7)
    _draw_box(ax, 5.0, 0.6, 2.0, 0.55, "Orders + OrderItems", "#F3E8FF", fs=7)
    _draw_box(ax, 8.5, 0.6, 1.8, 0.55, "DailySalesSummaries", "#C6F6D5", fs=7)
    _arrow(ax, 2.5, 1.6, 4.0, 1.6)
    _arrow(ax, 6.0, 1.85, 7.6, 2.1)
    _arrow(ax, 6.0, 1.35, 6.0, 0.88)
    _arrow(ax, 6.0, 0.6, 7.6, 0.6)
    ax.set_title("Lean Insights Pipeline (No ProductRecommendations / ProductPerformanceMetrics)", fontsize=11,
                 fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def make_evolution_diagram():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(10, 3.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.5)
    ax.axis("off")
    _draw_box(ax, 1.8, 1.8, 2.2, 1.2, "Full Design\n81 Tables", "#FED7D7", fs=9)
    _draw_box(ax, 5.0, 1.8, 2.2, 1.2, "V1 Draft\n47 Tables", "#FEEBC8", fs=9)
    _draw_box(ax, 8.2, 1.8, 2.2, 1.2, "V1 FINAL\n43 Tables", "#C6F6D5", fs=9)
    _arrow(ax, 2.9, 1.8, 3.9, 1.8)
    _arrow(ax, 6.1, 1.8, 7.1, 1.8)
    ax.set_title("Architecture Evolution", fontsize=12, fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def make_order_flow():
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(11, 4))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 4)
    ax.axis("off")
    cod = ["Pending", "Processing", "Shipped", "Delivered", "Completed"]
    for i, s in enumerate(cod):
        _draw_box(ax, 0.8 + i * 2.0, 3.0, 1.5, 0.5, s, "#D4EDDA", fs=7)
        if i < len(cod) - 1:
            _arrow(ax, 0.8 + i * 2.0 + 0.75, 3.0, 0.8 + (i + 1) * 2.0 - 0.75, 3.0)
    ax.text(5.5, 3.65, "Cash On Delivery", ha="center", fontsize=9, fontweight="bold", color="#276749")
    wallet = ["Pending", "Awaiting\nPayment", "Proof\nSubmitted", "Confirmed", "Processing", "Shipped", "Delivered"]
    for i, s in enumerate(wallet):
        _draw_box(ax, 0.5 + i * 1.45, 1.15, 1.15, 0.55, s, "#FFF3E0", fs=6)
        if i < len(wallet) - 1:
            _arrow(ax, 0.5 + i * 1.45 + 0.58, 1.15, 0.5 + (i + 1) * 1.45 - 0.58, 1.15)
    ax.text(5.5, 1.85, "Wallet + PaymentProofImages + Admin Approval", ha="center", fontsize=8,
            fontweight="bold", color="#C05621")
    ax.set_title("Order & Payment Workflows", fontsize=12, fontweight="bold", color="#1A365D")
    return _save_fig(fig)


def add_diagram(doc, title, png_bytes):
    add_para(doc, title, bold=True, size=12)
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.2))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Figure: {title}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()


def build_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Cairo Bags")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("V1 FINAL Database Architecture")
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = s2.add_run("43 Tables  |  Single Premium Bag Store  |  CODRA Development")
    r3.font.size = Pt(13)
    r3.italic = True
    s3 = doc.add_paragraph()
    s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = s3.add_run(datetime.now().strftime("%B %d, %Y"))
    r4.font.size = Pt(11)
    doc.add_page_break()

    add_heading(doc, "Table of Contents", 1)
    for item in [
        "1. Final Design Decisions",
        "2. Architecture Evolution (81 → 47 → 43)",
        "3. Final Module Structure",
        "4. Localization & Currency (without reference tables)",
        "5. Lean Insights Model",
        "6. Updated ERD & Diagrams",
        "7. Final V1 Table List (43 tables)",
        "8. Final Table Count by Category",
        "9. Tables Removed in This Revision",
        "10. Why Each Table Exists",
        "11. Payment & Shipping Reference",
        "12. Background Jobs",
    ]:
        add_para(doc, item)
    doc.add_page_break()

    # 1. Decisions
    add_heading(doc, "1. Final Design Decisions", 1)
    add_table(doc, ["#", "Decision", "Action"],
              [
                  ["1", "Languages table", "REMOVED — LanguageCode CHAR(2) enum: en, ar"],
                  ["2", "Currencies table", "REMOVED — EGP only, CurrencyCode on Orders"],
                  ["3", "ProductRecommendations", "REMOVED — use TrendingProducts + queries"],
                  ["4", "ProductPerformanceMetrics", "REMOVED — use DailySalesSummaries + OrderItems"],
                  ["5", "Governorates", "KEPT"],
                  ["6", "UserProductViews", "KEPT"],
                  ["7", "TrendingProducts", "KEPT"],
                  ["8", "DailySalesSummaries", "KEPT"],
                  ["9", "ReviewImages", "KEPT"],
                  ["10", "PaymentProofImages", "KEPT"],
              ])

    # 2. Evolution
    add_heading(doc, "2. Architecture Evolution", 1)
    add_diagram(doc, "Architecture Evolution", make_evolution_diagram())
    add_table(doc, ["Version", "Tables", "Notes"],
              [["Full architecture", "81", "Enterprise / multi-warehouse baseline"],
               ["V1 draft", "47", "Single-store simplification"],
               ["V1 FINAL", "43", "Lean reference tables + lean insights"]])

    # 3. Module structure
    add_heading(doc, "3. Final Module Structure", 1)
    add_diagram(doc, "Final Module Structure", make_final_module_diagram())
    add_table(doc, ["Module", "Tables", "Count"],
              [
                  ["Identity & Access", "AspNetUsers (+6 Identity), CustomerProfiles, PasswordResetOtps", "9"],
                  ["Reference", "Governorates", "1"],
                  ["Catalog", "Categories, CategoryTranslations, Products, ProductTranslations, ProductVariants, ProductImages", "6"],
                  ["Inventory", "Inventory, InventoryMovements", "2"],
                  ["Cart & Wishlist", "Carts, CartItems, Wishlists, WishlistItems", "4"],
                  ["Orders", "Orders, OrderItems, OrderStatusHistories, ShippingAddresses", "4"],
                  ["Shipping", "ShippingZones, ShippingZoneTranslations", "2"],
                  ["Payments", "PaymentMethods, PaymentMethodTranslations, OrderPayments, PaymentProofImages", "4"],
                  ["Reviews", "ProductReviews, ReviewImages", "2"],
                  ["Coupons", "Coupons, CouponUsages", "2"],
                  ["Notifications", "Notifications", "1"],
                  ["Banners", "Banners", "1"],
                  ["Analytics", "DailySalesSummaries", "1"],
                  ["Trending / Behavior", "UserProductViews, TrendingProducts", "2"],
                  ["System", "SystemSettings, AuditLogs", "2"],
                  ["TOTAL", "", "43"],
              ])

    # 4. Localization
    add_heading(doc, "4. Localization & Currency", 1)
    add_heading(doc, "4.1 Language (no Languages table)", 2)
    add_table(doc, ["Item", "Approach"],
              [["Storage", "LanguageCode CHAR(2) — values: en, ar"],
               ["Translation tables", "CategoryTranslations, ProductTranslations, ShippingZoneTranslations, PaymentMethodTranslations"],
               ["Inline bilingual", "Banners, OrderItems snapshots"],
               ["User default", "AspNetUsers.PreferredLanguage"],
               ["Store default", "SystemSettings.DefaultLanguageCode"]])
    add_heading(doc, "4.2 Currency (no Currencies table)", 2)
    add_table(doc, ["Item", "Approach"],
              [["Store currency", "EGP only for V1"],
               ["Orders", "CurrencyCode CHAR(3) DEFAULT 'EGP'"],
               ["SystemSettings", "DefaultCurrencyCode CHAR(3) DEFAULT 'EGP'"],
               ["Display", "ج.م / EGP in frontend"]])

    # 5. Insights
    add_heading(doc, "5. Lean Insights Model", 1)
    add_para(doc, "Removed ProductRecommendations and ProductPerformanceMetrics. V1 uses a minimal pipeline:")
    add_diagram(doc, "Lean Insights Pipeline", make_insights_pipeline_diagram())
    add_table(doc, ["Need", "V1 Solution"],
              [["Homepage trending", "TrendingProducts (computed from UserProductViews + sales)"],
               ["Admin revenue dashboard", "DailySalesSummaries"],
               ["Best sellers", "Query OrderItems directly"],
               ["Cross-sell on product page", "Same-category query or TrendingProducts — no extra table"]])

    # 6. ERD
    add_heading(doc, "6. Updated ERD & Diagrams", 1)
    add_diagram(doc, "V1 FINAL Core ERD", make_final_erd_diagram())
    add_diagram(doc, "Order & Payment Workflows", make_order_flow())

    doc.add_page_break()

    # 7. Full table list
    add_heading(doc, "7. Final V1 Table List — 43 Tables", 1)
    final_tables = [
        (1, "AspNetUsers", "Identity", "Auth, PreferredLanguage en/ar"),
        (2, "AspNetRoles", "Identity", "Admin, Customer"),
        (3, "AspNetUserRoles", "Identity", ""),
        (4, "AspNetUserClaims", "Identity", ""),
        (5, "AspNetRoleClaims", "Identity", ""),
        (6, "AspNetUserLogins", "Identity", "Google OAuth"),
        (7, "AspNetUserTokens", "Identity", ""),
        (8, "CustomerProfiles", "Customer", "Profile extension"),
        (9, "PasswordResetOtps", "Auth", "OTP reset"),
        (10, "Governorates", "Reference", "27 governorates → shipping zone"),
        (11, "Categories", "Catalog", "Hierarchy"),
        (12, "CategoryTranslations", "Catalog", "Name, Slug — LanguageCode"),
        (13, "Products", "Catalog", "Status, featured flags"),
        (14, "ProductTranslations", "Catalog", "Name, Description, Slug"),
        (15, "ProductVariants", "Catalog", "ColorNameAr/En, Sku, Price"),
        (16, "ProductImages", "Catalog", "Gallery + per-color"),
        (17, "Inventory", "Inventory", "1:1 per variant"),
        (18, "InventoryMovements", "Inventory", "Stock audit trail"),
        (19, "Carts", "Commerce", "User or guest"),
        (20, "CartItems", "Commerce", "Variant + quantity"),
        (21, "Wishlists", "Commerce", "One per user"),
        (22, "WishlistItems", "Commerce", "Saved products"),
        (23, "Orders", "Orders", "Snapshot address, CurrencyCode=EGP"),
        (24, "OrderItems", "Orders", "Immutable line snapshots"),
        (25, "OrderStatusHistories", "Orders", "Status audit"),
        (26, "ShippingAddresses", "Orders", "Address book"),
        (27, "ShippingZones", "Shipping", "Cairo / Giza / Other"),
        (28, "ShippingZoneTranslations", "Shipping", "Bilingual names"),
        (29, "PaymentMethods", "Payments", "6 Egyptian methods"),
        (30, "PaymentMethodTranslations", "Payments", "Bilingual instructions"),
        (31, "OrderPayments", "Payments", "Status + admin review"),
        (32, "PaymentProofImages", "Payments", "Transfer screenshots"),
        (33, "ProductReviews", "Reviews", "Rating + moderation"),
        (34, "ReviewImages", "Reviews", "Customer photos"),
        (35, "Coupons", "Marketing", "Discount codes"),
        (36, "CouponUsages", "Marketing", "Usage tracking"),
        (37, "Notifications", "Notifications", "In-app inbox"),
        (38, "Banners", "Marketing", "Inline bilingual"),
        (39, "DailySalesSummaries", "Analytics", "Daily revenue KPIs"),
        (40, "UserProductViews", "Trending", "Browse events"),
        (41, "TrendingProducts", "Trending", "Precomputed ranks"),
        (42, "SystemSettings", "System", "Store config, EGP, default lang"),
        (43, "AuditLogs", "System", "Admin change log"),
    ]
    add_table(doc, ["#", "Table", "Domain", "Notes"], final_tables)

    # 8. Count
    add_heading(doc, "8. Final Table Count by Category", 1)
    add_table(doc, ["Category", "Count"],
              [["ASP.NET Identity", "7"], ["Auth & Customer", "2"], ["Reference", "1"],
               ["Catalog", "6"], ["Inventory", "2"], ["Cart & Wishlist", "4"],
               ["Orders & Addresses", "4"], ["Shipping", "2"], ["Payments", "4"],
               ["Reviews", "2"], ["Coupons", "2"], ["Notifications", "1"],
               ["Banners", "1"], ["Analytics", "1"], ["Trending / Behavior", "2"],
               ["System", "2"], ["TOTAL V1 FINAL", "43"]])

    # 9. Removed
    add_heading(doc, "9. Tables Removed in This Revision", 1)
    add_table(doc, ["Removed Table", "Reason"],
              [["Languages", "Only 2 languages — CHAR(2) enum sufficient"],
               ["Currencies", "EGP-only store — no multi-currency V1"],
               ["ProductRecommendations", "TrendingProducts + same-category queries"],
               ["ProductPerformanceMetrics", "DailySalesSummaries + OrderItems queries"]])
    add_para(doc, "Previous V1 draft: 47 tables  →  V1 FINAL: 43 tables  (−4)", bold=True)

    # 10. Why
    add_heading(doc, "10. Why Each Table Exists", 1)
    why = [
        ("AspNetUsers (+ Identity)", "Real customer and admin authentication"),
        ("CustomerProfiles", "Profile separate from login credentials"),
        ("PasswordResetOtps", "Secure password recovery"),
        ("Governorates", "Egyptian addresses mapped to shipping zones"),
        ("Categories + Translations", "Bag taxonomy, bilingual SEO"),
        ("Products + Translations", "Catalog with ar/en content"),
        ("ProductVariants", "Sellable unit: color + SKU + price"),
        ("ProductImages", "Premium gallery and per-color images"),
        ("Inventory", "Single-store stock, prevent overselling"),
        ("InventoryMovements", "Accountability for every stock change"),
        ("Carts + CartItems", "Shopping before checkout"),
        ("Wishlists + WishlistItems", "Save bags for later"),
        ("Orders + OrderItems", "Purchases with immutable snapshots"),
        ("OrderStatusHistories", "Tracking and dispute resolution"),
        ("ShippingAddresses", "Repeat customer addresses"),
        ("ShippingZones + Translations", "Cairo / Giza / Other pricing"),
        ("PaymentMethods + Translations", "6 Egyptian payment methods"),
        ("OrderPayments + PaymentProofImages", "Wallet verification workflow"),
        ("ProductReviews + ReviewImages", "Social proof with photos"),
        ("Coupons + CouponUsages", "Discount codes with limits"),
        ("Notifications", "Order and payment alerts"),
        ("Banners", "Storefront marketing"),
        ("DailySalesSummaries", "Admin revenue dashboard"),
        ("UserProductViews", "Browse data for trending computation"),
        ("TrendingProducts", "Fast homepage trending list"),
        ("SystemSettings", "Store configuration"),
        ("AuditLogs", "Admin accountability"),
    ]
    add_table(doc, ["Table(s)", "Why"], why)

    # 11. Payment & Shipping
    add_heading(doc, "11. Payment & Shipping Reference", 1)
    add_table(doc, ["Zone", "Code", "Governorates"],
              [["Cairo", "CAIRO", "Cairo"],
               ["Giza", "GIZA", "Giza"],
               ["Other Governorates", "OTHER", "All remaining 25 governorates"]])
    add_table(doc, ["#", "Method", "Proof", "Admin Approval"],
              [["1", "Cash On Delivery", "No", "No"],
               ["2", "InstaPay", "Yes", "Yes"],
               ["3", "Vodafone Cash", "Yes", "Yes"],
               ["4", "Orange Cash", "Yes", "Yes"],
               ["5", "Etisalat Cash", "Yes", "Yes"],
               ["6", "WE Pay", "Yes", "Yes"]])

    # 12. Jobs
    add_heading(doc, "12. Background Jobs (V1 FINAL)", 1)
    add_table(doc, ["Job", "Input", "Output"],
              [["Aggregate daily sales", "Orders, OrderItems", "DailySalesSummaries"],
               ["Compute trending", "UserProductViews, OrderItems", "TrendingProducts"],
               ["Expire stock reservations", "Inventory.QuantityReserved", "Release held stock"],
               ["Low stock alert", "Inventory", "Notifications"],
               ["Payment review queue", "OrderPayments (ProofSubmitted)", "Admin notification"]])

    doc.add_page_break()
    add_heading(doc, "Document Information", 1)
    add_para(doc, "Cairo Bags V1 FINAL Database Architecture — CODRA Development")
    add_para(doc, "Architecture documentation only. No code, entities, DbContext, or migrations.")
    add_para(doc, f"File: {OUTPUT}")
    add_para(doc, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    return doc


if __name__ == "__main__":
    print("Building Cairo Bags V1 FINAL report...")
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Size: {os.path.getsize(OUTPUT):,} bytes")
